#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Translit Lite - Document Maker
Professional editor with Native DOCX, PDF Export, Print Layout, and Phonetic Typing.
"""

from __future__ import annotations

import json
import logging
import os
import re
import sys
import tempfile
from datetime import datetime
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional, Set, Tuple

from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtGui import QIcon, QTextCharFormat, QFont, QKeySequence
from PySide6.QtCore import Qt, QTimer
from PySide6.QtPrintSupport import QPrinter, QPrintPreviewDialog, QPageSetupDialog

# Configure Logging for Professional Debugging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Dependencies Check ---
try:
    from indic_transliteration.sanscript import transliterate  # type: ignore
    HAS_TRANSLIT = True
except ImportError:
    transliterate = None
    HAS_TRANSLIT = False
    logging.warning("indic-transliteration not installed.")

try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except ImportError:
    docx = None
    HAS_DOCX = False
    logging.warning("python-docx not installed.")

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    pd = None
    HAS_PANDAS = False
    logging.warning("pandas not installed.")

try:
    import speech_recognition as sr  # type: ignore
    HAS_SR = True
except ImportError:
    sr = None
    HAS_SR = False
    logging.warning("SpeechRecognition not installed.")

# --- Constants & Defaults ---
APP_NAME = "Translit Lite"
USER_DIR = Path(
    QtCore.QStandardPaths.writableLocation(QtCore.QStandardPaths.StandardLocation.AppDataLocation)
    or Path.home() / f".{APP_NAME.lower().replace(' ', '_')}"
)
USER_DIR.mkdir(parents=True, exist_ok=True)

USER_DICT_PATH = USER_DIR / "user_translit.json"
PHRASES_PATH = USER_DIR / "phrases.json"
SUGGESTIONS_PATH = USER_DIR / "suggestions.json"
TEMPLATES_PATH = USER_DIR / "templates.json"

DEFAULT_PHRASES = [
    "नमस्ते,\nआशा है आप स्वस्थ हैं।",
    "धन्यवाद,\n[आपका नाम]",
    "सादर,\n[आपका नाम]\n[पद]",
]

DEFAULT_DICT = {
    "yahan": "यहाँ", "yah": "यह", "rha": "रहा", "kya": "क्या",
    "main": "मैं", "hoon": "हूँ", "aur": "और", "ki": "कि", "kee": "की",
    "hota": "होता", "hoti": "होती", "hote": "होते", "bahut": "बहुत",
    "Ri": "ऋ", "R": "ृ", "O": "ॉ", "M": "ँ", "H": "ः", "^": "ं",
    "Gy": "ज्ञ्", "Ksh": "क्ष्", "Tr": "त्र्", "Shr": "श्र्",
    "Dhy": "ध्य्", "Dy": "द्य्", "Tt": "ट्ट्", "Dd": "ड्ड्"
}

U_LINE = "<p align='center'>__________________________________________________________________________________</p>"

# --- Data Model ---
@dataclass
class AppState:
    transliteration_enabled: bool = True
    font_family: str = "Nirmala UI"
    font_size: int = 16
    user_dict: Dict[str, str] = field(default_factory=lambda: DEFAULT_DICT.copy())
    phrases: List[str] = field(default_factory=lambda: DEFAULT_PHRASES.copy())
    suggestion_words: Set[str] = field(default_factory=set)
    templates: Dict[str, str] = field(default_factory=dict)
    active_dict_name: str = "Default"
    active_sugg_name: str = "Default"

    def load(self):
        self.user_dict = self._load_json(USER_DICT_PATH, self.user_dict)
        self.phrases = self._load_json(PHRASES_PATH, self.phrases)
        sugg_list = self._load_json(SUGGESTIONS_PATH, list(self.suggestion_words))
        self.suggestion_words = set(sugg_list)
        self.templates = self._load_json(TEMPLATES_PATH, {"Blank Document": "<p></p>"})

    def _load_json(self, path: Path, default: any) -> any:
        if path.exists():
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    return data if data else default
            except json.JSONDecodeError as e:
                logging.error(f"Corrupted JSON in {path}: {e}")
        else:
            self._save_json(path, default)
        return default

    def _save_json(self, path: Path, data: any) -> bool:
        try:
            with open(path, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
            return True
        except IOError as e:
            logging.error(f"Failed to save {path}: {e}")
            return False

    def save_all(self):
        self._save_json(USER_DICT_PATH, self.user_dict)
        self._save_json(PHRASES_PATH, self.phrases)
        self._save_json(SUGGESTIONS_PATH, list(self.suggestion_words))
        self._save_json(TEMPLATES_PATH, self.templates)

# --- Core Logic ---
class AdaptiveTransliterator:
    def __init__(self, state: AppState):
        if not HAS_TRANSLIT:
            raise RuntimeError("Missing indic-transliteration.")
        self.src, self.dst = "itrans", "devanagari"
        self.state = state
        self._reverse_cache = {}
        self._cached_dict_keys = []
        self._dict_hash = 0
        self._update_cache()

    def _update_cache(self):
        current_hash = hash(frozenset(self.state.user_dict.items()))
        if current_hash != self._dict_hash:
            self._cached_dict_keys = sorted(self.state.user_dict.keys(), key=len, reverse=True)
            self._dict_hash = current_hash

    def translit_token(self, latin: str) -> str:
        if not latin: return ""
        if latin in self.state.user_dict: return self.state.user_dict[latin]
        if latin.lower() in self.state.user_dict: return self.state.user_dict[latin.lower()]
            
        self._update_cache()
        temp_latin = latin
        
        for k in self._cached_dict_keys:
            is_modifier = (not k.isalpha()) or any(c.isupper() for c in k)
            if is_modifier and k in temp_latin:
                dev_val = self.state.user_dict[k]
                if dev_val not in self._reverse_cache:
                    try:
                        self._reverse_cache[dev_val] = transliterate(dev_val, self.dst, self.src)
                    except Exception:
                        self._reverse_cache[dev_val] = dev_val
                temp_latin = temp_latin.replace(k, self._reverse_cache[dev_val])
                
        force_halant = temp_latin.endswith('~')
        if force_halant: temp_latin = temp_latin[:-1]
                
        try: 
            res = transliterate(temp_latin, self.src, self.dst)
            if not force_halant and res.endswith('्'): res = res[:-1]
            return res
        except Exception: 
            return temp_latin

    def translit_full(self, text: str) -> str:
        return " ".join([self.translit_token(w) for w in (text or "").split()])

# --- Exporters ---
def export_native_to_docx(qdoc: QtGui.QTextDocument, filepath: str):
    if not HAS_DOCX: raise Exception("python-docx is not installed.")
    
    doc = docx.Document()
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Pt(0)
    
    def process_paragraph(block, docx_parent=None):
        if block.text().strip() == "" and not block.begin().fragment().charFormat().isImageFormat():
            if docx_parent is not None: docx_parent.add_paragraph()
            return

        p = docx_parent.add_paragraph() if docx_parent is not None else doc.add_paragraph()
        align = block.blockFormat().alignment()
        if align & Qt.AlignmentFlag.AlignHCenter: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align & Qt.AlignmentFlag.AlignRight: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align & Qt.AlignmentFlag.AlignJustify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        it = block.begin()
        while not it.atEnd():
            fragment = it.fragment()
            if fragment.isValid():
                fmt = fragment.charFormat()
                if fmt.isImageFormat():
                    img_name = fmt.toImageFormat().name()
                    img = qdoc.resource(QtGui.QTextDocument.ResourceType.ImageResource, QtCore.QUrl(img_name))
                    if img and not img.isNull():
                        fd, tmp = tempfile.mkstemp(suffix=".png")
                        os.close(fd)
                        try:
                            img.save(tmp, "PNG")
                            w = fmt.toImageFormat().width()
                            if w > 0:
                                run = p.add_run()
                                run.add_picture(tmp, width=Inches(w / 96.0))
                        finally:
                            if os.path.exists(tmp): os.remove(tmp) # Memory leak fixed
                else:
                    txt = fragment.text().replace('\u2028', '\n')
                    if txt:
                        run = p.add_run(txt)
                        font = fmt.font()
                        run.font.bold = font.bold() or (fmt.fontWeight() >= QFont.Weight.Bold)
                        run.font.italic = font.italic() or fmt.fontItalic()
                        run.font.underline = font.underline() or fmt.fontUnderline()
                        
                        f_size = font.pointSize()
                        if f_size > 0: run.font.size = Pt(f_size)
                        
                        fg_brush = fmt.foreground()
                        if fg_brush.style() != Qt.BrushStyle.NoBrush:
                            c = fg_brush.color()
                            if c.isValid() and c.alpha() > 0 and c.name().lower() != "#000000":
                                run.font.color.rgb = RGBColor(c.red(), c.green(), c.blue())
            it += 1

    for i in range(qdoc.blockCount()):
        process_paragraph(qdoc.findBlockByNumber(i), doc)

    try: 
        doc.save(filepath)
    except PermissionError: 
        raise PermissionError(f"Close the document {filepath} in Word before saving.")

# --- Threads ---
class SpeechWorker(QtCore.QThread):
    partial = QtCore.Signal(str)
    finished_text = QtCore.Signal(str)
    error = QtCore.Signal(str)

    def __init__(self, lang: str = "hi-IN", parent=None):
        super().__init__(parent)
        self._stop = False
        self.lang = lang
        self._accum: List[str] = []

    def stop(self):
        self._stop = True

    def run(self):
        if not HAS_SR:
            self.error.emit("SpeechRecognition missing.")
            return
        try:
            rec = sr.Recognizer()
            with sr.Microphone() as source:
                rec.adjust_for_ambient_noise(source, duration=0.3)
                while not self._stop:
                    try: 
                        audio = rec.listen(source, timeout=0.5, phrase_time_limit=5)
                        text = rec.recognize_google(audio, language=self.lang)
                        self._accum.append(text)
                        self.partial.emit(text)
                    except sr.WaitTimeoutError:
                        continue # Safely loops allowing _stop to interrupt
                    except Exception: 
                        continue
        except Exception as e: 
            self.error.emit(str(e))
        finally:
            self.finished_text.emit(" ".join(self._accum).strip())

# --- Editor Component ---
class TranslitEditor(QtWidgets.QTextEdit):
    countsChanged = QtCore.Signal(int, int)
    cursorPositionChangedDetailed = QtCore.Signal(int, int)
    prefixChanged = QtCore.Signal(str)
    insertSuggestionTrigger = QtCore.Signal()

    def __init__(self, translit: AdaptiveTransliterator, state: AppState, parent=None):
        super().__init__(parent)
        self.translit = translit
        self.state = state
        self.setAcceptRichText(True)
        self.setFont(QtGui.QFont(self.state.font_family, self.state.font_size))
        self.setPlaceholderText("Start typing phonetically. Press Ctrl+L to dictate.")

        self._composing_latin = ""
        self._composing_start_pos: Optional[int] = None
        self._composing_display_len = 0
        self._ignore_cursor_move = False
        self._has_suggestions = False
        
        self.textChanged.connect(self._emit_counts)
        self.cursorPositionChanged.connect(self._on_cursor_moved)

    def keyPressEvent(self, ev: QtGui.QKeyEvent):
        key, mods = ev.key(), ev.modifiers()
        
        if key == Qt.Key.Key_Tab and self._has_suggestions:
            self.insertSuggestionTrigger.emit()
            return
            
        if mods & Qt.KeyboardModifier.ControlModifier: 
            return super().keyPressEvent(ev)

        if key in {Qt.Key.Key_Left, Qt.Key.Key_Right, Qt.Key.Key_Up, Qt.Key.Key_Down, Qt.Key.Key_Return, Qt.Key.Key_Enter}:
            if self._composing_latin: self._commit_composing()
            return super().keyPressEvent(ev)

        if key == Qt.Key.Key_Backspace and self._composing_latin:
            self._composing_latin = self._composing_latin[:-1]
            if self._composing_latin: 
                self._update_composing()
            else: 
                self._remove_composing()
            self.prefixChanged.emit(self.translit.translit_token(self._composing_latin))
            return
        
        txt = ev.text()
        if txt.isprintable() and len(txt) == 1 and not txt.isspace() and ord(txt) < 128:
            self._handle_typing(txt)
            return

        if self._composing_latin: self._commit_composing()
        return super().keyPressEvent(ev)

    def _handle_typing(self, txt: str):
        cur = self.textCursor()
        if not self._composing_latin:
            if cur.hasSelection(): cur.removeSelectedText()
            self._composing_start_pos = cur.position()
            self._composing_latin = txt
            translit_text = self.translit.translit_token(self._composing_latin) if self.state.transliteration_enabled else self._composing_latin
            
            self._ignore_cursor_move = True
            cur.insertText(translit_text, cur.charFormat())
            self._composing_display_len = len(translit_text)
            self.setTextCursor(cur)
            self._ignore_cursor_move = False
        else:
            self._composing_latin += txt
            translit_text = self.translit.translit_token(self._composing_latin) if self.state.transliteration_enabled else self._composing_latin
            self._update_composing_with_text(translit_text)
            
        self.prefixChanged.emit(translit_text)

    def insert_suggestion(self, word: str):
        cur = self.textCursor()
        if self._composing_latin: self._remove_composing()
        else:
            cur.movePosition(QtGui.QTextCursor.MoveOperation.StartOfWord, QtGui.QTextCursor.MoveMode.KeepAnchor)
            cur.removeSelectedText()
        cur.insertText(word + " ")
        self.setTextCursor(cur)
        self.setFocus()
        self._has_suggestions = False
        self.prefixChanged.emit("")

    def _update_composing_with_text(self, text: str):
        if self._composing_start_pos is None: return
        self._ignore_cursor_move = True
        cur = self.textCursor()
        cur.setPosition(self._composing_start_pos)
        cur.setPosition(self._composing_start_pos + self._composing_display_len, QtGui.QTextCursor.MoveMode.KeepAnchor)
        cur.insertText(text, cur.charFormat())
        self._composing_display_len = len(text)
        self.setTextCursor(cur)
        self._ignore_cursor_move = False

    def _update_composing(self):
        text = self.translit.translit_token(self._composing_latin) if self.state.transliteration_enabled else self._composing_latin
        self._update_composing_with_text(text)

    def _remove_composing(self):
        if self._composing_start_pos is None: return
        self._ignore_cursor_move = True
        cur = self.textCursor()
        cur.setPosition(self._composing_start_pos)
        cur.setPosition(self._composing_start_pos + self._composing_display_len, QtGui.QTextCursor.MoveMode.KeepAnchor)
        cur.removeSelectedText()
        self._composing_latin, self._composing_start_pos, self._composing_display_len = "", None, 0
        self.setTextCursor(cur)
        self._ignore_cursor_move = False

    def _commit_composing(self):
        if not self._composing_latin: return
        self._composing_latin, self._composing_start_pos, self._composing_display_len = "", None, 0
        self._emit_counts()

    def _on_cursor_moved(self):
        if self._ignore_cursor_move: return
        if self._composing_latin: self._commit_composing()
        
        cur = self.textCursor()
        cur.movePosition(QtGui.QTextCursor.MoveOperation.StartOfWord, QtGui.QTextCursor.MoveMode.KeepAnchor)
        self.prefixChanged.emit(cur.selectedText().strip())
        self.cursorPositionChangedDetailed.emit(cur.blockNumber() + 1, cur.positionInBlock())

    def _emit_counts(self):
        txt = self.toPlainText()
        words = len([w for w in re.split(r"\s+", txt.strip()) if w]) if txt.strip() else 0
        self.countsChanged.emit(words, len(txt))

# --- Main Window ---
class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(APP_NAME)
        self.resize(1100, 750)
        
        self.state = AppState()
        self.state.load()
        self.current_filepath: Optional[str] = None
        self.is_dirty = False
        self.printer = QPrinter(QPrinter.PrinterMode.HighResolution)

        try: 
            self.translit = AdaptiveTransliterator(self.state)
        except RuntimeError as e:
            QtWidgets.QMessageBox.critical(self, "Dependency Error", str(e))
            sys.exit(1)

        self.editor = TranslitEditor(self.translit, self.state, self)
        self.setCentralWidget(self.editor)
        
        self.editor.textChanged.connect(lambda: setattr(self, 'is_dirty', True))
        self.editor.countsChanged.connect(self._update_counts)
        self.editor.cursorPositionChangedDetailed.connect(lambda l, c: self.lbl_cursor.setText(f"Ln {l}, Col {c}"))
        self.editor.prefixChanged.connect(self._update_suggestions)

        self.sr_worker: Optional[SpeechWorker] = None
        QtWidgets.QApplication.instance().installEventFilter(self)

        self._setup_ui()
        self._new_file()

    def _setup_ui(self):
        # Toolbars
        self.tb_file = self.addToolBar("File")
        self.tb_fmt = self.addToolBar("Format")
        self.addToolBarBreak()

        # Actions
        def add_action(tb, name, slot, shortcut=None):
            act = QtGui.QAction(name, self)
            if shortcut: act.setShortcut(shortcut)
            act.triggered.connect(slot)
            tb.addAction(act)
            return act

        add_action(self.tb_file, "New", self._new_file, "Ctrl+N")
        add_action(self.tb_file, "Open", self._open_file, "Ctrl+O")
        add_action(self.tb_file, "Save", self._save_file, "Ctrl+S")
        add_action(self.tb_file, "Export PDF", self._export_pdf)
        self.tb_file.addSeparator()
        
        add_action(self.tb_fmt, "B", lambda: self._format('bold'), "Ctrl+B")
        add_action(self.tb_fmt, "I", lambda: self._format('italic'), "Ctrl+I")
        add_action(self.tb_fmt, "U", lambda: self._format('underline'), "Ctrl+U")
        self.tb_fmt.addSeparator()

        self.font_combo = QtWidgets.QFontComboBox()
        self.font_combo.setCurrentFont(QtGui.QFont(self.state.font_family))
        self.font_combo.currentFontChanged.connect(lambda f: self._format('font', f.family()))
        self.tb_fmt.addWidget(self.font_combo)

        self.size_spin = QtWidgets.QSpinBox()
        self.size_spin.setRange(8, 72); self.size_spin.setValue(self.state.font_size)
        self.size_spin.valueChanged.connect(lambda s: self._format('size', s))
        self.tb_fmt.addWidget(self.size_spin)

        # Status Bar
        self.status = self.statusBar()
        self.lbl_cursor = QtWidgets.QLabel("Ln 1, Col 0")
        self.lbl_counts = QtWidgets.QLabel("0 words • 0 chars")
        self.lbl_speech = QtWidgets.QLabel("🎤 Hold Ctrl+L to dictation")
        
        for lbl in (self.lbl_cursor, self.lbl_counts, self.lbl_speech):
            lbl.setStyleSheet("padding: 0 10px;")
            self.status.addPermanentWidget(lbl)

        # Sidebars
        self._build_suggestions_dock()

    def _build_suggestions_dock(self):
        self.dock = QtWidgets.QDockWidget("Suggestions (Tab to insert)", self)
        self.sugg_list = QtWidgets.QListWidget()
        self.sugg_list.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self.sugg_list.itemClicked.connect(lambda i: self.editor.insert_suggestion(i.text()))
        self.dock.setWidget(self.sugg_list)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.dock)

    def _format(self, fmt_type, val=None):
        cur = self.editor.textCursor()
        fmt = cur.charFormat()
        
        if fmt_type == 'bold': fmt.setFontWeight(QtGui.QFont.Weight.Bold if fmt.fontWeight() != QtGui.QFont.Weight.Bold else QtGui.QFont.Weight.Normal)
        elif fmt_type == 'italic': fmt.setFontItalic(not fmt.fontItalic())
        elif fmt_type == 'underline': fmt.setFontUnderline(not fmt.fontUnderline())
        elif fmt_type == 'font': fmt.setFontFamily(val)
        elif fmt_type == 'size': fmt.setFontPointSize(val)
            
        if cur.hasSelection(): cur.mergeCharFormat(fmt)
        else: self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()

    def eventFilter(self, obj, event):
        if event.type() == QtCore.QEvent.Type.KeyPress:
            if event.key() == Qt.Key.Key_L and (event.modifiers() & Qt.KeyboardModifier.ControlModifier) and not event.isAutoRepeat():
                self._start_speech()
                return True
        elif event.type() == QtCore.QEvent.Type.KeyRelease:
            if event.key() == Qt.Key.Key_L and not event.isAutoRepeat():
                self._stop_speech()
                return True
        return super().eventFilter(obj, event)

    def _start_speech(self):
        if self.sr_worker and self.sr_worker.isRunning(): return
        self.sr_worker = SpeechWorker(parent=self)
        self.sr_worker.partial.connect(lambda s: self.lbl_speech.setText(f"Listening: {s}..."))
        self.sr_worker.finished_text.connect(self._on_speech_done)
        self.sr_worker.start()
        self.lbl_speech.setText("Listening...")

    def _stop_speech(self):
        if self.sr_worker:
            self.sr_worker.stop()
            self.lbl_speech.setText("Processing audio...")

    def _on_speech_done(self, text):
        if text:
            try: dev_text = self.translit.translit_full(text)
            except Exception: dev_text = text
            self.editor.textCursor().insertText(dev_text + " ")
        self.lbl_speech.setText("🎤 Hold Ctrl+L to dictation")
        self.sr_worker = None

    def _update_suggestions(self, prefix):
        self.sugg_list.clear()
        if not prefix: 
            self.editor._has_suggestions = False
            return
            
        matches = [w for w in self.state.suggestion_words if w.lower().startswith(prefix.lower())]
        self.sugg_list.addItems(sorted(matches, key=len)[:50])
        
        if self.sugg_list.count() > 0:
            self.sugg_list.setCurrentRow(0)
            self.editor._has_suggestions = True
        else:
            self.editor._has_suggestions = False

    def _update_counts(self, w, c):
        self.lbl_counts.setText(f"{w} words • {c} chars")

    def _new_file(self):
        if self.is_dirty:
            res = QtWidgets.QMessageBox.question(self, "Save?", "Save current document?", 
                                                 QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No)
            if res == QtWidgets.QMessageBox.StandardButton.Yes: self._save_file()
        self.editor.clear()
        self.current_filepath = None
        self.is_dirty = False

    def _open_file(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Open", "", "Documents (*.docx *.txt *.html);;All Files (*.*)")
        if not path: return
        ext = path.lower().split('.')[-1]
        
        try:
            if ext == 'txt':
                with open(path, 'r', encoding='utf-8') as f: self.editor.setPlainText(f.read())
            elif ext == 'html':
                with open(path, 'r', encoding='utf-8') as f: self.editor.setHtml(f.read())
            elif ext == 'docx' and HAS_DOCX:
                doc = docx.Document(path)
                self.editor.setPlainText("\n".join([p.text for p in doc.paragraphs])) # Fallback to plain read for speed
            
            self.current_filepath = path
            self.is_dirty = False
            self.status.showMessage(f"Opened {path}", 3000)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Error", str(e))

    def _save_file(self):
        if not self.current_filepath:
            path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save", "", "Word Document (*.docx);;HTML File (*.html)")
            if not path: return
            self.current_filepath = path
            
        ext = self.current_filepath.lower().split('.')[-1]
        try:
            if ext == 'docx' and HAS_DOCX:
                export_native_to_docx(self.editor.document(), self.current_filepath)
            else:
                with open(self.current_filepath, 'w', encoding='utf-8') as f:
                    f.write(self.editor.toHtml() if ext == 'html' else self.editor.toPlainText())
            self.is_dirty = False
            self.status.showMessage("Saved successfully.", 3000)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Save Error", str(e))

    def _export_pdf(self):
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Export PDF", "", "PDF Document (*.pdf)")
        if path:
            self.printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
            self.printer.setOutputFileName(path)
            self.editor.document().print_(self.printer)
            self.status.showMessage("Exported to PDF.", 3000)

    def closeEvent(self, event):
        self.state.save_all()
        if self.is_dirty:
            res = QtWidgets.QMessageBox.question(self, "Save before exit?", "You have unsaved changes. Save?", 
                                                 QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No)
            if res == QtWidgets.QMessageBox.StandardButton.Yes: self._save_file()
        event.accept()

if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    app.setStyle("Fusion") # Cleaner modern look across platforms
    
    # Modern base font fallback
    available = QtGui.QFontDatabase.families()
    target_font = next((f for f in ["Nirmala UI", "Mangal", "Arial"] if f in available), "Arial")
    app.setFont(QtGui.QFont(target_font, 10))
    
    win = MainWindow()
    win.show()
    sys.exit(app.exec())