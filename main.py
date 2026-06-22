#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Translit - Document Maker
Professional, high-performance editor with Native DOCX, PDF Export, Print Layout, Page Setup, Advanced Templates, and AI Typing Predictions.
"""

from __future__ import annotations

import json
import logging
import os
import re
import sys
import tempfile
import difflib
import math
from datetime import datetime
from dataclasses import dataclass, field
from pathlib import Path

from PySide6 import QtCore, QtGui, QtWidgets
from PySide6.QtGui import QIcon, QTextCharFormat, QFont
from PySide6.QtCore import Qt, QTimer
from PySide6.QtPrintSupport import QPrinter, QPrintPreviewDialog, QPageSetupDialog

import urllib.request
import urllib.parse

# --- LOCAL MODULE IMPORTS ---
import templates
import style
import splash

logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')

# Dependencies
try:
    from indic_transliteration.sanscript import transliterate  # type: ignore
    HAS_TRANSLIT = True
except Exception:
    transliterate = None
    HAS_TRANSLIT = False

try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_COLOR_INDEX, WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    HAS_DOCX = True
except Exception:
    docx = None
    HAS_DOCX = False

try:
    import pandas as pd
    HAS_PANDAS = True
except Exception:
    pd = None
    HAS_PANDAS = False

try:
    import speech_recognition as sr  # type: ignore
    HAS_SR = True
except Exception:
    sr = None
    HAS_SR = False

try:
    import pypdf
    HAS_PDF = True
except Exception:
    pypdf = None
    HAS_PDF = False

APP_NAME = "Translit"
USER_DIR = Path(
    QtCore.QStandardPaths.writableLocation(QtCore.QStandardPaths.StandardLocation.AppDataLocation)
    or Path.home() / ".hindi_office_wysiwyg"
)
USER_DIR.mkdir(parents=True, exist_ok=True)
USER_DICT_PATH = USER_DIR / "user_translit.json"
PHRASES_PATH = USER_DIR / "phrases.json"
SUGGESTIONS_PATH = USER_DIR / "suggestions.json"
TEMPLATES_PATH = USER_DIR / "templates.json"
NEXT_WORDS_PATH = USER_DIR / "next_words.json"

DEFAULT_PHRASES = [
    "नमस्ते,\nआशा है आप स्वस्थ हैं।",
    "धन्यवाद,\n[आपका नाम]",
    "सादर,\n[आपका नाम]\n[पद]",
]

DEFAULT_DICT = {
    "yahan": "यहाँ", "yah": "यह", "rha": "रहा", "narkhi": "नारखी",
    "kya": "क्या", "koko": "कोमल", "karne": "करने", "hi": "हाय",
    "hemlata": "हेमलता", "hello": "हेलो", "han": "हाँ", "etah": "एटा",
    "dr.": "डॉ.", "aj": "आज", "main": "मैं", "hoon": "हूँ", "aur": "और",
    "ki": "कि", "kee": "की", "hota": "होता", "hoti": "होती", "hote": "होते",
    "bahut": "बहुत", "Ri": "ऋ", "R": "ृ", "O": "ॉ", "M": "ँ", "H": "ः",
    "^": "ं", "*": ".", "!": "ॠ", "Aum": "ॐ", "Z": "़",
    "Gy": "ज्ञ्", "Ksh": "क्ष्", "Tr": "त्र्", "Shr": "श्र्",
    "Dhy": "ध्य्", "Dy": "द्य्", "Tt": "ट्ट्", "Dd": "ड्ड्", "Tth": "ट्ठ्",
    "Ru": "रु", "Roo": "रू"
}

def save_json_atomically(data: dict | list, filepath: Path) -> bool:
    tmp_path = filepath.with_suffix('.tmp')
    try:
        with open(tmp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        os.replace(tmp_path, filepath) 
        return True
    except Exception as e:
        logging.error(f"Failed to save {filepath.name}: {e}")
        if tmp_path.exists():
            tmp_path.unlink()
        return False


@dataclass
class AppState:
    transliteration_enabled: bool = True
    font_family: str = "Nirmala UI"
    font_size: int = 16
    user_dict: dict[str, str] = field(default_factory=lambda: DEFAULT_DICT.copy())
    phrases: list[str] = field(default_factory=lambda: DEFAULT_PHRASES.copy())
    suggestion_words: set[str] = field(default_factory=set)
    next_words: dict[str, dict[str, int]] = field(default_factory=dict)
    
    templates: dict[str, str] = field(default_factory=lambda: {
        "Question Paper": templates.TPL_QUESTION_PAPER,
        "Application Form": templates.TPL_APPLICATION,
        "Declaration Letter": templates.TPL_DECLARATION
    })
    
    active_dict_name: str = "Default (user_translit.json)"
    active_sugg_name: str = "Default (suggestions.json)"
    active_pred_name: str = "Default (next_words.json)"
    
    view_mode: str = "Portrait"
    suggestion_mode: str = "Inline"
    dict_is_dirty: bool = True

    def load(self):
        if USER_DICT_PATH.exists():
            try:
                with open(USER_DICT_PATH, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, dict) and data: 
                        self.user_dict = data
                        self.dict_is_dirty = True
            except Exception as e: logging.error(f"Error loading dict: {e}")
        else:
            self.save_user_dict() 

        if PHRASES_PATH.exists():
            try:
                with open(PHRASES_PATH, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list): self.phrases = data
            except Exception as e: logging.error(f"Error loading phrases: {e}")
            
        if SUGGESTIONS_PATH.exists():
            try:
                with open(SUGGESTIONS_PATH, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, list): self.suggestion_words = set(data)
            except Exception as e: logging.error(f"Error loading suggestions: {e}")

        if TEMPLATES_PATH.exists():
            try:
                with open(TEMPLATES_PATH, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, dict) and data: self.templates = data
            except Exception as e: logging.error(f"Error loading templates: {e}")
        else:
            self.save_templates()

        if NEXT_WORDS_PATH.exists():
            try:
                with open(NEXT_WORDS_PATH, "r", encoding="utf-8") as f:
                    self.next_words = json.load(f)
            except Exception as e: logging.error(f"Error loading predictions: {e}")
        else:
            self.scan_templates_for_predictions()
            self.save_next_words()

    def learn_bigram(self, w1: str, w2: str):
        if not w1 or not w2: return
        w1, w2 = w1.lower(), w2.lower()
        if w1 not in self.next_words: self.next_words[w1] = {}
        self.next_words[w1][w2] = self.next_words[w1].get(w2, 0) + 1

    def scan_templates_for_predictions(self):
        for name, html in self.templates.items():
            text = re.sub('<[^<]+>', ' ', html)
            words = [w for w in re.findall(r'[a-zA-Z]+|[\u0900-\u097F]+', text) if len(w) > 1]
            for i in range(len(words) - 1):
                self.learn_bigram(words[i], words[i+1])

    def save_next_words(self, path: Path | None = None) -> bool:
        if len(self.next_words) > 15000:
            sorted_items = sorted(
                self.next_words.items(), 
                key=lambda x: sum(x[1].values()), 
                reverse=True
            )
            self.next_words = dict(sorted_items[:15000])
        return save_json_atomically(self.next_words, path or NEXT_WORDS_PATH)

    def save_user_dict(self, path: Path | None = None) -> bool:
        return save_json_atomically(self.user_dict, path or USER_DICT_PATH)

    def save_phrases(self):
        return save_json_atomically(self.phrases, PHRASES_PATH)

    def save_suggestions(self, path: Path | None = None) -> bool:
        return save_json_atomically(list(self.suggestion_words), path or SUGGESTIONS_PATH)

    def save_templates(self, path: Path | None = None) -> bool:
        return save_json_atomically(self.templates, path or TEMPLATES_PATH)


class AdaptiveTransliterator:
    def __init__(self, state: AppState):
        if not HAS_TRANSLIT:
            raise RuntimeError("Missing indic-transliteration.")
        self.src = "itrans"
        self.dst = "devanagari"
        self.state = state
        self._reverse_cache = {}
        self._cached_dict_keys = []
        self._modifier_map = {}
        self._modifier_pattern = None
        self._update_cache()

    def _get_itrans(self, dev_str: str) -> str:      
        if dev_str not in self._reverse_cache:
            try:
                self._reverse_cache[dev_str] = transliterate(dev_str, self.dst, self.src)
            except Exception:
                self._reverse_cache[dev_str] = dev_str
        return self._reverse_cache[dev_str]

    def _update_cache(self):
        if not self.state.dict_is_dirty: return
        
        self._cached_dict_keys = sorted(self.state.user_dict.keys(), key=len, reverse=True)
        self._modifier_map = {}
        for k in self._cached_dict_keys:
            is_modifier = (not k.isalpha()) or any(c.isupper() for c in k)
            if is_modifier:
                dev_val = self.state.user_dict[k]
                self._modifier_map[k] = self._get_itrans(dev_val)
                
        if self._modifier_map:
            sorted_map_keys = sorted(self._modifier_map.keys(), key=len, reverse=True)
            self._modifier_pattern = re.compile('|'.join(map(re.escape, sorted_map_keys)))
        else:
            self._modifier_pattern = None
            
        self.state.dict_is_dirty = False

    def translit_token(self, latin: str) -> str:
        if not latin: return ""
        if latin in self.state.user_dict: return self.state.user_dict[latin]
        if latin.lower() in self.state.user_dict: return self.state.user_dict[latin.lower()]
            
        self._update_cache()
        
        if self._modifier_pattern:
            temp_latin = self._modifier_pattern.sub(lambda m: self._modifier_map[m.group(0)], latin)
        else:
            temp_latin = latin
                
        force_halant = temp_latin.endswith('~')
        if force_halant:
            temp_latin = temp_latin[:-1]
                
        try: 
            res = transliterate(temp_latin, self.src, self.dst)
            if not force_halant and res.endswith('्'):
                res = res[:-1]
            return res
        except Exception: 
            return temp_latin

    def translit_full(self, text: str) -> str:
        if not text: return ""
        words = text.split()
        return " ".join([self.translit_token(w) for w in words])


class DocxExportThread(QtCore.QThread):
    progress = QtCore.Signal(str)
    success = QtCore.Signal(str)
    error = QtCore.Signal(str)
    
    def __init__(self, doc_clone: QtGui.QTextDocument, filepath: str, parent=None):
        super().__init__(parent)
        self.doc_clone = doc_clone
        self.filepath = filepath
        
    def run(self):
        try:
            self.progress.emit("Formatting document structure...")
            export_native_to_docx(self.doc_clone, self.filepath, progress_callback=self.progress.emit)
            self.success.emit(self.filepath)
        except Exception as e:
            self.error.emit(str(e))

def export_native_to_docx(qdoc: QtGui.QTextDocument, filepath: str, progress_callback=None):
    if not HAS_DOCX: raise Exception("python-docx is not installed.")
    
    doc = docx.Document()
    style = doc.styles['Normal']
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.0
    
    visited_frames = set()

    def process_paragraph(block, docx_parent=None, p=None):
        if block.text().strip() == "" and not block.begin().fragment().charFormat().isImageFormat():
            if p is None and docx_parent is not None:
                docx_parent.add_paragraph()
            return

        if p is None and docx_parent is not None:
            p = docx_parent.add_paragraph()
            
        if p is None: return

        align = block.blockFormat().alignment()
        if align & Qt.AlignmentFlag.AlignHCenter: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align & Qt.AlignmentFlag.AlignRight: p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif align & Qt.AlignmentFlag.AlignJustify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        it = block.begin()
        while not it.atEnd():
            fragment = it.fragment()
            if fragment.isValid():
                fmt = fragment.charFormat()
                if fmt.isImageFormat():
                    img_fmt = fmt.toImageFormat()
                    img_name = img_fmt.name()
                    img = qdoc.resource(QtGui.QTextDocument.ResourceType.ImageResource, QtCore.QUrl(img_name))
                    if img and not img.isNull():
                        fd, tmp = tempfile.mkstemp(suffix=".png")
                        os.close(fd) 
                        try:
                            img.save(tmp, "PNG")
                            w = img_fmt.width()
                            if w > 0:
                                run = p.add_run()
                                run.add_picture(tmp, width=Inches(w / 96.0))
                        finally:
                            if os.path.exists(tmp): os.remove(tmp)
                else:
                    txt = fragment.text().replace('\u2028', '\n')
                    if txt:
                        run = p.add_run(txt)
                        font = fmt.font()
                        run.font.bold = font.bold() or (fmt.fontWeight() >= QFont.Weight.Bold)
                        run.font.italic = font.italic() or fmt.fontItalic()
                        run.font.underline = font.underline() or fmt.fontUnderline()
                        run.font.strike = font.strikeOut() or fmt.fontStrikeOut()
                        
                        if fmt.verticalAlignment() == QTextCharFormat.VerticalAlignment.AlignSubScript:
                            run.font.subscript = True
                        elif fmt.verticalAlignment() == QTextCharFormat.VerticalAlignment.AlignSuperScript:
                            run.font.superscript = True
                        
                        f_size = font.pointSize()
                        if f_size <= 0:
                            p_size = font.pixelSize()
                            if p_size > 0: f_size = p_size * 0.75
                        if f_size > 0: run.font.size = Pt(f_size)
                        
                        fg_brush = fmt.foreground()
                        if fg_brush.style() != Qt.BrushStyle.NoBrush:
                            c = fg_brush.color()
                            if c.isValid() and c.alpha() > 0 and c.name().lower() != "#000000":
                                run.font.color.rgb = RGBColor(c.red(), c.green(), c.blue())
                        
                        bg_brush = fmt.background()
                        if bg_brush.style() != Qt.BrushStyle.NoBrush:
                            c = bg_brush.color()
                            if c.isValid() and c.alpha() > 0 and c.name().lower() != "#ffffff":
                                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            it += 1

    def process_table(table: QtGui.QTextTable, docx_parent):
        if table.rows() == 0 or table.columns() == 0: return 
        
        docx_table = docx_parent.add_table(rows=table.rows(), cols=table.columns())
        docx_table.style = 'Table Grid'
        docx_table.autofit = False
        try:
            total_width = Inches(6.5)
            col_width = total_width / table.columns()
            for col in docx_table.columns: col.width = col_width
            for row in docx_table.rows:
                for cell in row.cells: cell.width = col_width
        except Exception: pass
        
        if WD_TABLE_ALIGNMENT is not None:
            t_align = table.format().alignment()
            if t_align & Qt.AlignmentFlag.AlignHCenter: docx_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            elif t_align & Qt.AlignmentFlag.AlignRight: docx_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            else: docx_table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for r in range(table.rows()):
            for c in range(table.columns()):
                cell = table.cellAt(r, c)
                docx_cell = docx_table.cell(r, c)
                is_first_para = True
                it = cell.begin()
                while not it.atEnd():
                    child_frame = it.currentFrame()
                    child_block = it.currentBlock()
                    if child_frame:
                        if isinstance(child_frame, QtGui.QTextTable): process_table(child_frame, docx_cell)
                        else: process_frame(child_frame, docx_cell)
                    elif child_block.isValid(): 
                        if is_first_para and len(docx_cell.paragraphs) > 0:
                            process_paragraph(child_block, docx_parent=None, p=docx_cell.paragraphs[0])
                            is_first_para = False
                        else:
                            process_paragraph(child_block, docx_parent=docx_cell)
                    it += 1

    def process_frame(frame, docx_parent):
        if id(frame) in visited_frames: return 
        visited_frames.add(id(frame))
        
        it = frame.begin()
        while not it.atEnd():
            child_frame = it.currentFrame()
            child_block = it.currentBlock()
            if child_frame:
                if isinstance(child_frame, QtGui.QTextTable): process_table(child_frame, docx_parent)
                else: process_frame(child_frame, docx_parent)
            elif child_block.isValid(): 
                process_paragraph(child_block, docx_parent=docx_parent)
            it += 1

    process_frame(qdoc.rootFrame(), doc)
    if progress_callback: progress_callback("Saving to disk...")
    try: doc.save(filepath)
    except PermissionError: raise PermissionError(f"Cannot save to {filepath}. Please close the document if it is open in another program.")


class TemplateManagerDialog(QtWidgets.QDialog):
    def __init__(self, state: AppState, editor_ref: QtWidgets.QTextEdit, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Advanced Template Manager")
        self.resize(900, 600)
        self.state = state
        self.editor_ref = editor_ref

        layout = QtWidgets.QVBoxLayout(self)
        main_h = QtWidgets.QHBoxLayout()
        left_v = QtWidgets.QVBoxLayout()
        self.list_widget = QtWidgets.QListWidget()
        self.list_widget.currentTextChanged.connect(self._on_select)
        left_v.addWidget(QtWidgets.QLabel("<b>Your Templates</b>"))
        left_v.addWidget(self.list_widget)
        
        self.btn_insert = QtWidgets.QPushButton("Insert to Editor")
        self.btn_insert.setStyleSheet("background-color: #10b981; color: white; font-weight:bold;")
        self.btn_add_current = QtWidgets.QPushButton("Save Editor as Template")
        self.btn_delete = QtWidgets.QPushButton("Delete Selected")
        left_v.addWidget(self.btn_insert)
        left_v.addWidget(self.btn_add_current)
        left_v.addWidget(self.btn_delete)
        main_h.addLayout(left_v, 1)

        right_v = QtWidgets.QVBoxLayout()
        right_v.addWidget(QtWidgets.QLabel("<b>Preview / Edit Template HTML</b>"))
        self.preview_editor = QtWidgets.QTextEdit()
        self.preview_editor.setAcceptRichText(True)
        right_v.addWidget(self.preview_editor)
        self.btn_save_edit = QtWidgets.QPushButton("Update Template")
        right_v.addWidget(self.btn_save_edit)
        main_h.addLayout(right_v, 3)

        layout.addLayout(main_h)
        bot_h = QtWidgets.QHBoxLayout()
        self.btn_export = QtWidgets.QPushButton("Export Group to File...")
        self.btn_import = QtWidgets.QPushButton("Import Group from File...")
        self.btn_close = QtWidgets.QPushButton("Close")
        bot_h.addWidget(self.btn_export)
        bot_h.addWidget(self.btn_import)
        bot_h.addStretch()
        bot_h.addWidget(self.btn_close)
        layout.addLayout(bot_h)

        self.btn_insert.clicked.connect(self._insert_template)
        self.btn_add_current.clicked.connect(self._add_from_editor)
        self.btn_delete.clicked.connect(self._delete_template)
        self.btn_save_edit.clicked.connect(self._update_template)
        self.btn_export.clicked.connect(self._export_group)
        self.btn_import.clicked.connect(self._import_group)
        self.btn_close.clicked.connect(self.accept)
        self._refresh_list()

    def _refresh_list(self):
        self.list_widget.clear()
        self.list_widget.addItems(sorted(self.state.templates.keys()))

    def _on_select(self, name: str):
        if name and name in self.state.templates:
            self.preview_editor.setHtml(self.state.templates[name])

    def _insert_template(self):
        name = self.list_widget.currentItem().text() if self.list_widget.currentItem() else None
        if name:
            if hasattr(self.editor_ref, "_commit_composing"): self.editor_ref._commit_composing()
            self.editor_ref.insertHtml(self.state.templates[name])
            self.accept()

    def _add_from_editor(self):
        name, ok = QtWidgets.QInputDialog.getText(self, "Save Template", "Template Name:")
        if ok and name.strip():
            self.state.templates[name.strip()] = self.editor_ref.toHtml()
            self.state.save_templates()
            self._refresh_list()
            self.list_widget.setCurrentItem(self.list_widget.findItems(name.strip(), Qt.MatchFlag.MatchExactly)[0])
            QtWidgets.QMessageBox.information(self, "Success", "Template saved!")

    def _delete_template(self):
        item = self.list_widget.currentItem()
        if item:
            name = item.text()
            ans = QtWidgets.QMessageBox.question(self, "Delete", f"Delete template '{name}'?")
            if ans == QtWidgets.QMessageBox.StandardButton.Yes:
                del self.state.templates[name]
                self.state.save_templates()
                self._refresh_list()
                self.preview_editor.clear()

    def _update_template(self):
        item = self.list_widget.currentItem()
        if item:
            name = item.text()
            self.state.templates[name] = self.preview_editor.toHtml()
            self.state.save_templates()
            QtWidgets.QMessageBox.information(self, "Success", f"Template '{name}' updated.")

    def _export_group(self):
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Export Template Group", "", "JSON Files (*.json)")
        if path:
            try:
                with open(path, "w", encoding="utf-8") as f:
                    json.dump(self.state.templates, f, ensure_ascii=False, indent=2)
                QtWidgets.QMessageBox.information(self, "Exported", f"Templates exported to {path}")
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Error", str(e))

    def _import_group(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Import Template Group", "", "JSON Files (*.json)")
        if path:
            try:
                with open(path, "r", encoding="utf-8") as f:
                    data = json.load(f)
                    if isinstance(data, dict):
                        self.state.templates.update(data)
                        self.state.save_templates()
                        self._refresh_list()
                        QtWidgets.QMessageBox.information(self, "Imported", "Templates loaded successfully!")
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Error", str(e))


class SpeechWorker(QtCore.QThread):
    partial = QtCore.Signal(str)
    finished_text = QtCore.Signal(str)
    error = QtCore.Signal(str)

    def __init__(self, lang: str = "hi-IN", parent=None):
        super().__init__(parent)
        self._stop = False
        self.lang = lang
        self._accum: list[str] = []

    def stop(self): self._stop = True

    def run(self):
        if not HAS_SR:
            self.error.emit("SpeechRecognition not installed.")
            return
        try:
            rec = sr.Recognizer()
            with sr.Microphone() as source:
                try: rec.adjust_for_ambient_noise(source, duration=0.4)
                except Exception: pass
                while not self._stop:
                    try: 
                        audio = rec.listen(source, timeout=0.5, phrase_time_limit=5)
                    except sr.WaitTimeoutError:
                        continue
                    except Exception:
                        if self._stop: break
                        continue
                    try:
                        text = rec.recognize_google(audio, language=self.lang)
                        self._accum.append(text)
                        self.partial.emit(text)
                    except Exception: continue
        except Exception as e: self.error.emit(str(e))
        finally:
            self.finished_text.emit(" ".join(self._accum).strip())

class FileScannerThread(QtCore.QThread):
    progress = QtCore.Signal(str)
    finished_words = QtCore.Signal(set)
    finished_bigrams = QtCore.Signal(list)
    error = QtCore.Signal(str)

    def __init__(self, filepaths: list[str], mode: str = "words", parent=None):
        super().__init__(parent)
        self.filepaths = filepaths
        self.mode = mode

    def run(self):
        extracted_words = set()
        extracted_bigrams = []
        try:
            for path in self.filepaths:
                self.progress.emit(f"Scanning {os.path.basename(path)}...")
                ext = path.lower().split('.')[-1]
                full_text = ""
                
                if ext == "txt":
                    with open(path, "r", encoding="utf-8") as f: full_text = f.read()
                elif ext == "docx" and HAS_DOCX:
                    doc = docx.Document(path)
                    full_text = " ".join([p.text for p in doc.paragraphs])
                elif ext == "pdf" and HAS_PDF:
                    with open(path, "rb") as f:
                        reader = pypdf.PdfReader(f)
                        full_text = " ".join([page.extract_text() for page in reader.pages if page.extract_text()])
                elif ext in ["csv", "xlsx", "ods"]:
                    if HAS_PANDAS:
                        if ext == "xlsx": df = pd.read_excel(path, engine="openpyxl")
                        elif ext == "ods": df = pd.read_excel(path, engine="odf")
                        else: df = pd.read_csv(path)
                        df = df.fillna("") 
                        full_text = " ".join([str(val) for val in df.values.flatten() if str(val).strip()])

                if full_text:
                    words = [w for w in re.findall(r'[a-zA-Z]+|[\u0900-\u097F]+', full_text) if len(w) > 1]
                    if self.mode == "words":
                        extracted_words.update(words)
                    elif self.mode == "bigrams":
                        for i in range(len(words) - 1):
                            extracted_bigrams.append((words[i], words[i+1]))

            if self.mode == "words":
                self.finished_words.emit(extracted_words)
            else:
                self.finished_bigrams.emit(extracted_bigrams)
                
        except Exception as e:
            self.error.emit(str(e))

class FuzzySearchThread(QtCore.QThread):
    result_ready = QtCore.Signal(list)
    def __init__(self, prefix_lower, all_words, max_results=5, cutoff=0.6, parent=None):
        super().__init__(parent)
        self.prefix_lower = prefix_lower
        self.all_words = all_words
        self.max_results = max_results
        self.cutoff = cutoff

    def run(self):
        candidates = [w for w in self.all_words if abs(len(w) - len(self.prefix_lower)) <= 2]
        fuzzies = difflib.get_close_matches(self.prefix_lower, candidates, n=self.max_results, cutoff=self.cutoff)
        matches = []
        for f in fuzzies:
            matched_orig = next((w for w in self.all_words if w.lower() == f.lower()), f)
            matches.append(matched_orig)
        self.result_ready.emit(matches)


class SuggestionsManagerDialog(QtWidgets.QDialog):
    def __init__(self, state: AppState, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Suggestions Database Manager")
        self.resize(750, 600)
        self.state = state

        layout = QtWidgets.QVBoxLayout(self)
        self.lbl_dict = QtWidgets.QLabel(f"<b>Active Database:</b> {self.state.active_sugg_name} | <b>Total Words:</b> {len(self.state.suggestion_words)}")
        self.lbl_dict.setStyleSheet("font-size: 14px; padding: 5px;")
        layout.addWidget(self.lbl_dict)

        self.table = QtWidgets.QTableWidget(0, 1)
        self.table.verticalHeader().setDefaultSectionSize(30)
        self.table.setHorizontalHeaderLabels(["Suggestion Word"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.setSortingEnabled(True)
        layout.addWidget(self.table)

        btns_top = QtWidgets.QHBoxLayout()
        self.btn_add = QtWidgets.QPushButton("Add Word")
        self.btn_remove = QtWidgets.QPushButton("Remove Selected")
        self.btn_apply = QtWidgets.QPushButton("Apply (runtime)")
        btns_top.addWidget(self.btn_add); btns_top.addWidget(self.btn_remove); btns_top.addWidget(self.btn_apply)
        layout.addLayout(btns_top)

        btns_mid = QtWidgets.QHBoxLayout()
        self.btn_load_def = QtWidgets.QPushButton("Load Default")
        self.btn_save_def = QtWidgets.QPushButton("Save Default")
        self.btn_load_man = QtWidgets.QPushButton("Load Manual File...")
        self.btn_save_man = QtWidgets.QPushButton("Save Manual File...")
        btns_mid.addWidget(self.btn_load_def); btns_mid.addWidget(self.btn_save_def)
        btns_mid.addWidget(self.btn_load_man); btns_mid.addWidget(self.btn_save_man)
        layout.addLayout(btns_mid)

        btns_bot = QtWidgets.QHBoxLayout()
        self.btn_scan = QtWidgets.QPushButton("Scan/Extract Words from File(s)...")
        self.btn_scan.setStyleSheet("background-color: #10b981; color: white; font-weight: bold;")
        self.btn_close = QtWidgets.QPushButton("Close")
        btns_bot.addWidget(self.btn_scan)
        btns_bot.addStretch()
        btns_bot.addWidget(self.btn_close)
        layout.addLayout(btns_bot)

        self.btn_add.clicked.connect(self.add_row)
        self.btn_remove.clicked.connect(self.remove_selected)
        self.btn_apply.clicked.connect(self.apply_runtime)
        self.btn_load_def.clicked.connect(self.load_default)
        self.btn_save_def.clicked.connect(lambda: self.save_to_file(SUGGESTIONS_PATH))
        self.btn_load_man.clicked.connect(self.load_manual)
        self.btn_save_man.clicked.connect(lambda: self.save_to_file(None))
        self.btn_scan.clicked.connect(self.scan_file)
        self.btn_close.clicked.connect(self.accept)
        self.reload_table()

    def reload_table(self, data_set=None):
        self.table.setSortingEnabled(False)
        self.table.setRowCount(0)
        source = data_set if data_set is not None else self.state.suggestion_words
        sample = list(source)[:2000]
        for w in sample:
            r = self.table.rowCount()
            self.table.insertRow(r)
            self.table.setItem(r, 0, QtWidgets.QTableWidgetItem(w))
        self.table.setSortingEnabled(True)
        self.lbl_dict.setText(f"<b>Active Database:</b> {self.state.active_sugg_name} | <b>Total Words:</b> {len(source)}")

    def add_row(self):
            self.table.setSortingEnabled(False)
            self.table.insertRow(0) 
            self.table.setItem(0, 0, QtWidgets.QTableWidgetItem(""))
            self.table.scrollToItem(self.table.item(0, 0))
            self.table.editItem(self.table.item(0, 0))
            self.table.setSortingEnabled(True)

    def remove_selected(self):
        rows = sorted({i.row() for i in self.table.selectedIndexes()}, reverse=True)
        for r in rows: self.table.removeRow(r)

    def apply_runtime(self):
        new = set()
        for r in range(self.table.rowCount()):
            w = self.table.item(r, 0)
            if w and w.text().strip(): new.add(w.text().strip())
        self.state.suggestion_words = new
        self.state.active_sugg_name = "Runtime / Unsaved"
        self.lbl_dict.setText(f"<b>Active Database:</b> {self.state.active_sugg_name} | <b>Total Words:</b> {len(self.state.suggestion_words)}")
        QtWidgets.QMessageBox.information(self, "Applied", "Suggestions applied to runtime memory.")

    def save_to_file(self, target_path=None):
        if not target_path:
            target_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Manual Suggestions", "", "JSON Files (*.json)")
            if not target_path: return
        self.apply_runtime()
        try:
            with open(target_path, "w", encoding="utf-8") as f:
                json.dump(list(self.state.suggestion_words), f, ensure_ascii=False, indent=2)
            self.state.active_sugg_name = os.path.basename(str(target_path))
            self.lbl_dict.setText(f"<b>Active Database:</b> {self.state.active_sugg_name} | <b>Total Words:</b> {len(self.state.suggestion_words)}")
            QtWidgets.QMessageBox.information(self, "Saved", f"Saved successfully to: {target_path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Save failed", str(e))

    def load_default(self):
        if SUGGESTIONS_PATH.exists():
            try:
                with open(SUGGESTIONS_PATH, "r", encoding="utf-8") as f: data = json.load(f)
                if isinstance(data, list):
                    self.state.suggestion_words = set(data)
                    self.state.active_sugg_name = "Default (suggestions.json)"
                    self.reload_table()
            except Exception as e: QtWidgets.QMessageBox.critical(self, "Load failed", str(e))

    def load_manual(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Load Manual Suggestions", "", "JSON Files (*.json)")
        if path:
            try:
                with open(path, "r", encoding="utf-8") as f: data = json.load(f)
                if isinstance(data, list):
                    self.state.suggestion_words = set(data)
                    self.state.active_sugg_name = os.path.basename(path)
                    self.reload_table()
            except Exception as e: QtWidgets.QMessageBox.critical(self, "Error", str(e))

    def scan_file(self):
        filters = "Documents (*.txt *.csv *.xlsx *.ods *.docx *.pdf);;All Files (*.*)"
        paths, _ = QtWidgets.QFileDialog.getOpenFileNames(self, "Scan Words from Files", "", filters)
        if not paths: return
        
        self.btn_scan.setEnabled(False)
        self.scanner = FileScannerThread(paths, mode="words", parent=self)
        self.scanner.progress.connect(lambda msg: self.btn_scan.setText(msg))
        self.scanner.finished_words.connect(self._on_scan_finished)
        self.scanner.error.connect(lambda err: QtWidgets.QMessageBox.critical(self, "Scan Error", err))
        self.scanner.start()

    def _on_scan_finished(self, extracted_words: set):
        added = 0
        for w in extracted_words:
            if len(w) > 1 and w not in self.state.suggestion_words:
                self.state.suggestion_words.add(w)
                added += 1
        self.reload_table()
        self.btn_scan.setEnabled(True)
        self.btn_scan.setText("Scan/Extract Words from File(s)...")
        QtWidgets.QMessageBox.information(self, "Scan Complete", f"Extracted and added {added} NEW unique words.")

class PredictionManagerDialog(QtWidgets.QDialog):
    def __init__(self, state: AppState, parent=None):
        super().__init__(parent)
        self.setWindowTitle("AI Prediction Model Manager")
        self.resize(800, 600)
        self.state = state

        layout = QtWidgets.QVBoxLayout(self)
        
        self.lbl_info = QtWidgets.QLabel(f"<b>Active Model:</b> {self.state.active_pred_name} | <b>Total Connections:</b> {self._get_total_connections()}")
        self.lbl_info.setStyleSheet("font-size: 14px; padding: 5px;")
        layout.addWidget(self.lbl_info)

        self.table = QtWidgets.QTableWidget(0, 3)
        self.table.setHorizontalHeaderLabels(["Word 1 (Context)", "Word 2 (Prediction)", "Frequency"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.setSortingEnabled(True)
        layout.addWidget(self.table)

        btns_mid = QtWidgets.QHBoxLayout()
        self.btn_load_def = QtWidgets.QPushButton("Load Default Model")
        self.btn_save_def = QtWidgets.QPushButton("Save Default Model")
        self.btn_load_man = QtWidgets.QPushButton("Load Custom Model...")
        self.btn_save_man = QtWidgets.QPushButton("Save Custom Model...")
        self.btn_clear = QtWidgets.QPushButton("Clear/New Model")
        
        btns_mid.addWidget(self.btn_load_def); btns_mid.addWidget(self.btn_save_def)
        btns_mid.addWidget(self.btn_load_man); btns_mid.addWidget(self.btn_save_man)
        btns_mid.addWidget(self.btn_clear)
        layout.addLayout(btns_mid)

        btns_bot = QtWidgets.QHBoxLayout()
        self.btn_scan = QtWidgets.QPushButton("Scan/Extract Conversational Flow from File(s)...")
        self.btn_scan.setStyleSheet("background-color: #3b82f6; color: white; font-weight: bold; padding: 8px;")
        self.btn_close = QtWidgets.QPushButton("Close")
        
        btns_bot.addWidget(self.btn_scan)
        btns_bot.addStretch()
        btns_bot.addWidget(self.btn_close)
        layout.addLayout(btns_bot)

        self.btn_load_def.clicked.connect(self.load_default)
        self.btn_save_def.clicked.connect(lambda: self.save_to_file(NEXT_WORDS_PATH))
        self.btn_load_man.clicked.connect(self.load_manual)
        self.btn_save_man.clicked.connect(lambda: self.save_to_file(None))
        self.btn_clear.clicked.connect(self.clear_model)
        self.btn_scan.clicked.connect(self.scan_file)
        self.btn_close.clicked.connect(self.accept)
        
        self.reload_table()

    def _get_total_connections(self):
        return sum(len(targets) for targets in self.state.next_words.values())

    def reload_table(self):
        self.table.setSortingEnabled(False)
        self.table.setRowCount(0)
        
        flat_list = []
        for w1, targets in self.state.next_words.items():
            for w2, count in targets.items():
                flat_list.append((w1, w2, count))
                
        flat_list.sort(key=lambda x: x[2], reverse=True)
        
        for w1, w2, count in flat_list[:2000]:
            r = self.table.rowCount()
            self.table.insertRow(r)
            self.table.setItem(r, 0, QtWidgets.QTableWidgetItem(w1))
            self.table.setItem(r, 1, QtWidgets.QTableWidgetItem(w2))
            
            count_item = QtWidgets.QTableWidgetItem()
            count_item.setData(Qt.ItemDataRole.DisplayRole, count)
            self.table.setItem(r, 2, count_item)
            
        self.table.setSortingEnabled(True)
        self.lbl_info.setText(f"<b>Active Model:</b> {self.state.active_pred_name} | <b>Total Connections:</b> {self._get_total_connections()}")

    def clear_model(self):
        ans = QtWidgets.QMessageBox.question(self, "Clear Model", "Are you sure you want to clear all learned predictions? This will start a fresh memory model.")
        if ans == QtWidgets.QMessageBox.StandardButton.Yes:
            self.state.next_words.clear()
            self.state.active_pred_name = "Unsaved Blank Model"
            self.reload_table()

    def save_to_file(self, target_path=None):
        if not target_path:
            target_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Prediction Model", "", "JSON Files (*.json)")
            if not target_path: return
            
        if self.state.save_next_words(target_path):
            self.state.active_pred_name = os.path.basename(str(target_path))
            self.lbl_info.setText(f"<b>Active Model:</b> {self.state.active_pred_name} | <b>Total Connections:</b> {self._get_total_connections()}")
            QtWidgets.QMessageBox.information(self, "Saved", f"Model saved successfully to: {target_path}")
        else:
            QtWidgets.QMessageBox.critical(self, "Save failed", "Could not save the model.")

    def load_default(self):
        if NEXT_WORDS_PATH.exists():
            try:
                with open(NEXT_WORDS_PATH, "r", encoding="utf-8") as f: 
                    self.state.next_words = json.load(f)
                self.state.active_pred_name = "Default (next_words.json)"
                self.reload_table()
                QtWidgets.QMessageBox.information(self, "Loaded", "Default model loaded successfully.")
            except Exception as e: QtWidgets.QMessageBox.critical(self, "Load failed", str(e))
        else:
            QtWidgets.QMessageBox.warning(self, "Not found", "Default model file does not exist yet.")

    def load_manual(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Load Prediction Model", "", "JSON Files (*.json)")
        if path:
            try:
                with open(path, "r", encoding="utf-8") as f: 
                    data = json.load(f)
                if isinstance(data, dict):
                    self.state.next_words = data
                    self.state.active_pred_name = os.path.basename(path)
                    self.reload_table()
                    QtWidgets.QMessageBox.information(self, "Loaded", "Custom model loaded successfully.")
            except Exception as e: QtWidgets.QMessageBox.critical(self, "Error", str(e))

    def scan_file(self):
        filters = "Documents (*.txt *.csv *.xlsx *.ods *.docx *.pdf);;All Files (*.*)"
        paths, _ = QtWidgets.QFileDialog.getOpenFileNames(self, "Scan Context from Files", "", filters)
        if not paths: return
        
        self.btn_scan.setEnabled(False)
        self.scanner = FileScannerThread(paths, mode="bigrams", parent=self)
        self.scanner.progress.connect(lambda msg: self.btn_scan.setText(msg))
        self.scanner.finished_bigrams.connect(self._on_scan_finished)
        self.scanner.error.connect(lambda err: QtWidgets.QMessageBox.critical(self, "Scan Error", err))
        self.scanner.start()

    def _on_scan_finished(self, extracted_bigrams: list):
        for w1, w2 in extracted_bigrams:
            self.state.learn_bigram(w1, w2)
            
        self.reload_table()
        self.btn_scan.setEnabled(True)
        self.btn_scan.setText("Scan/Extract Conversational Flow from File(s)...")
        QtWidgets.QMessageBox.information(self, "Scan Complete", f"Successfully analyzed files and learned {len(extracted_bigrams)} context pairs.")


class CorrectionsDialog(QtWidgets.QDialog):
    def __init__(self, state: AppState, editor_ref: Optional[QtWidgets.QTextEdit] = None, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Dictionary Training & File Scanner")
        self.resize(950, 650)
        self.state = state
        self.editor_ref = editor_ref

        layout = QtWidgets.QVBoxLayout(self)
        self.lbl_dict = QtWidgets.QLabel(f"<b>Active Dictionary:</b> {self.state.active_dict_name}")
        self.lbl_dict.setStyleSheet("font-size: 14px; padding: 5px;")
        layout.addWidget(self.lbl_dict)

        self.table = QtWidgets.QTableWidget(0, 2)
        self.table.setHorizontalHeaderLabels(["Shortcut / Latin (key)", "Devanagari / Value (value)"])
        self.table.horizontalHeader().setStretchLastSection(True)
        self.table.setSortingEnabled(True)
        layout.addWidget(self.table)

        btns_top = QtWidgets.QHBoxLayout()
        self.btn_add = QtWidgets.QPushButton("Add Row")
        self.btn_remove = QtWidgets.QPushButton("Remove Selected")
        self.btn_apply = QtWidgets.QPushButton("Apply (runtime)")
        btns_top.addWidget(self.btn_add); btns_top.addWidget(self.btn_remove); btns_top.addWidget(self.btn_apply)
        layout.addLayout(btns_top)

        btns_mid = QtWidgets.QHBoxLayout()
        self.btn_load_def = QtWidgets.QPushButton("Load Default")
        self.btn_save_def = QtWidgets.QPushButton("Save Default")
        self.btn_load_man = QtWidgets.QPushButton("Load Manual File...")
        self.btn_save_man = QtWidgets.QPushButton("Save Manual File...")
        btns_mid.addWidget(self.btn_load_def); btns_mid.addWidget(self.btn_save_def)
        btns_mid.addWidget(self.btn_load_man); btns_mid.addWidget(self.btn_save_man)
        layout.addLayout(btns_mid)

        btns_bot = QtWidgets.QHBoxLayout()
        self.btn_scan = QtWidgets.QPushButton("Scan/Extract Words from File(s)...")
        self.btn_scan.setStyleSheet("background-color: #10b981; color: white; font-weight: bold;")
        self.btn_apply_sel = QtWidgets.QPushButton("Apply mapping to selection")
        self.btn_close = QtWidgets.QPushButton("Close")
        btns_bot.addWidget(self.btn_scan)
        btns_bot.addWidget(self.btn_apply_sel)
        btns_bot.addStretch()
        btns_bot.addWidget(self.btn_close)
        layout.addLayout(btns_bot)

        self.btn_add.clicked.connect(self.add_row)
        self.btn_remove.clicked.connect(self.remove_selected)
        self.btn_apply.clicked.connect(self.apply_runtime)
        self.btn_load_def.clicked.connect(self.load_default)
        self.btn_save_def.clicked.connect(lambda: self.save_to_file(USER_DICT_PATH))
        self.btn_load_man.clicked.connect(self.load_manual)
        self.btn_save_man.clicked.connect(lambda: self.save_to_file(None))
        self.btn_scan.clicked.connect(self.scan_file)
        self.btn_apply_sel.clicked.connect(self.apply_mapping_to_selection)
        self.btn_close.clicked.connect(self.accept)
        self.reload_table()

    def reload_table(self, data_dict=None):
        self.table.setSortingEnabled(False)
        self.table.setRowCount(0)
        source = data_dict if data_dict is not None else self.state.user_dict
        for k, v in source.items():
            r = self.table.rowCount()
            self.table.insertRow(r)
            self.table.setItem(r, 0, QtWidgets.QTableWidgetItem(k))
            self.table.setItem(r, 1, QtWidgets.QTableWidgetItem(v))
        self.table.setSortingEnabled(True)

    def add_row(self):
        self.table.setSortingEnabled(False)
        self.table.insertRow(0) 
        self.table.setItem(0, 0, QtWidgets.QTableWidgetItem(""))
        self.table.setItem(0, 1, QtWidgets.QTableWidgetItem(""))
        self.table.scrollToItem(self.table.item(0, 0))
        self.table.editItem(self.table.item(0, 0))
        self.table.setSortingEnabled(True)

    def remove_selected(self):
        rows = sorted({i.row() for i in self.table.selectedIndexes()}, reverse=True)
        for r in rows: self.table.removeRow(r)

    def apply_runtime(self):
        new = {}
        for r in range(self.table.rowCount()):
            k, v = self.table.item(r, 0), self.table.item(r, 1)
            if k and v and k.text().strip(): new[k.text().strip()] = v.text()
        self.state.user_dict = new
        self.state.dict_is_dirty = True
        self.state.active_dict_name = "Runtime / Unsaved"
        self.lbl_dict.setText(f"<b>Active Dictionary:</b> {self.state.active_dict_name}")
        QtWidgets.QMessageBox.information(self, "Applied", "Corrections applied to runtime memory.")

    def save_to_file(self, target_path=None):
        if not target_path:
            target_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Manual Dictionary", "", "JSON Files (*.json)")
            if not target_path: return
        new = {}
        for r in range(self.table.rowCount()):
            k, v = self.table.item(r, 0), self.table.item(r, 1)
            if k and v and k.text().strip(): new[k.text().strip()] = v.text()
        try:
            with open(target_path, "w", encoding="utf-8") as f:
                json.dump(new, f, ensure_ascii=False, indent=2)
            self.state.user_dict = new
            self.state.dict_is_dirty = True
            self.state.active_dict_name = os.path.basename(str(target_path))
            self.lbl_dict.setText(f"<b>Active Dictionary:</b> {self.state.active_dict_name}")
            QtWidgets.QMessageBox.information(self, "Saved", f"Saved successfully to: {target_path}")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Save failed", str(e))

    def load_default(self):
        if USER_DICT_PATH.exists():
            try:
                with open(USER_DICT_PATH, "r", encoding="utf-8") as f: data = json.load(f)
                if isinstance(data, dict):
                    self.state.user_dict = data
                    self.state.dict_is_dirty = True
                    self.state.active_dict_name = "Default (user_translit.json)"
                    self.lbl_dict.setText(f"<b>Active Dictionary:</b> {self.state.active_dict_name}")
                    self.reload_table()
            except Exception as e: QtWidgets.QMessageBox.critical(self, "Load failed", str(e))

    def load_manual(self):
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Load Manual Dictionary", "", "JSON Files (*.json)")
        if path:
            try:
                with open(path, "r", encoding="utf-8") as f: data = json.load(f)
                if isinstance(data, dict):
                    self.state.user_dict = data
                    self.state.dict_is_dirty = True
                    self.state.active_dict_name = os.path.basename(path)
                    self.lbl_dict.setText(f"<b>Active Dictionary:</b> {self.state.active_dict_name}")
                    self.reload_table()
            except Exception as e: QtWidgets.QMessageBox.critical(self, "Error", str(e))

    def scan_file(self):
        filters = "Documents (*.txt *.csv *.xlsx *.ods *.docx *.pdf);;All Files (*.*)"
        paths, _ = QtWidgets.QFileDialog.getOpenFileNames(self, "Scan Words from Files", "", filters)
        if not paths: return
        
        self.btn_scan.setEnabled(False)
        self.scanner = FileScannerThread(paths, mode="words", parent=self)
        self.scanner.progress.connect(lambda msg: self.btn_scan.setText(msg))
        self.scanner.finished_words.connect(self._on_scan_finished)
        self.scanner.error.connect(lambda err: QtWidgets.QMessageBox.critical(self, "Scan Error", err))
        self.scanner.start()

    def _on_scan_finished(self, extracted_words: set):
        existing_words = set()
        for r in range(self.table.rowCount()):
            item = self.table.item(r, 0)
            if item and item.text().strip(): existing_words.add(item.text().strip().lower())

        added = 0
        self.table.setSortingEnabled(False)
        for w in extracted_words:
            w_lower = w.lower()
            if w_lower not in existing_words and not w_lower.isnumeric():
                r = self.table.rowCount()
                self.table.insertRow(r)
                is_devanagari = bool(re.search(r"[\u0900-\u097F]", w_lower))
                if is_devanagari:
                    self.table.setItem(r, 0, QtWidgets.QTableWidgetItem(""))
                    self.table.setItem(r, 1, QtWidgets.QTableWidgetItem(w_lower))
                else:
                    self.table.setItem(r, 0, QtWidgets.QTableWidgetItem(w_lower))
                    self.table.setItem(r, 1, QtWidgets.QTableWidgetItem(""))
                existing_words.add(w_lower)
                added += 1
        self.table.setSortingEnabled(True)
        self.btn_scan.setEnabled(True)
        self.btn_scan.setText("Scan/Extract Words from File(s)...")
        QtWidgets.QMessageBox.information(self, "Scan Complete", f"Extracted and added {added} NEW unique words.")

    def apply_mapping_to_selection(self):
        if not self.editor_ref: return
        cur = self.editor_ref.textCursor()
        if not cur.hasSelection(): cur.select(QtGui.QTextCursor.SelectionType.WordUnderCursor)
        sel = cur.selectedText()
        if not sel: return
        mappings = [(self.table.item(r, 0).text().strip(), self.table.item(r, 1).text()) 
                    for r in range(self.table.rowCount()) if self.table.item(r, 0) and self.table.item(r, 1)]
        if not mappings: return
        dlg = QtWidgets.QDialog(self)
        dlg.setWindowTitle("Choose mapping")
        dlg.resize(480, 360)
        v = QtWidgets.QVBoxLayout(dlg)
        lw = QtWidgets.QListWidget(dlg)
        for k, v_ in mappings: lw.addItem(f"{k}  →  {v_}")
        v.addWidget(lw)
        h = QtWidgets.QHBoxLayout()
        b_apply = QtWidgets.QPushButton("Apply")
        b_apply.clicked.connect(lambda: self._do_apply(lw, mappings, cur, dlg))
        h.addWidget(b_apply); h.addWidget(QtWidgets.QPushButton("Cancel", clicked=dlg.reject))
        v.addLayout(h)
        dlg.exec()

    def _do_apply(self, lw, mappings, cur, dlg):
        idx = lw.currentRow()
        if idx >= 0:
            cur.insertText(mappings[idx][1])
            dlg.accept()


class FindReplaceDialog(QtWidgets.QDialog):
    def __init__(self, editor: QtWidgets.QTextEdit, parent=None):
        super().__init__(parent)
        self.editor = editor
        self.setWindowTitle("Find & Replace")
        self.resize(520, 180)
        v = QtWidgets.QVBoxLayout(self)
        form = QtWidgets.QFormLayout()
        self.find_edit = QtWidgets.QLineEdit()
        self.replace_edit = QtWidgets.QLineEdit()
        form.addRow("Find:", self.find_edit)
        form.addRow("Replace:", self.replace_edit)
        v.addLayout(form)
        opt = QtWidgets.QHBoxLayout()
        self.ck_case = QtWidgets.QCheckBox("Match case")
        self.ck_word = QtWidgets.QCheckBox("Whole word")
        opt.addWidget(self.ck_case); opt.addWidget(self.ck_word); opt.addStretch(1)
        v.addLayout(opt)
        h = QtWidgets.QHBoxLayout()
        self.btn_prev = QtWidgets.QPushButton("Find Prev")
        self.btn_next = QtWidgets.QPushButton("Find Next")
        self.btn_rep = QtWidgets.QPushButton("Replace")
        self.btn_all = QtWidgets.QPushButton("Replace All")
        self.btn_close = QtWidgets.QPushButton("Close")
        for b in (self.btn_prev, self.btn_next, self.btn_rep, self.btn_all, self.btn_close): h.addWidget(b)
        v.addLayout(h)
        self.btn_close.clicked.connect(self.reject)
        self.btn_prev.clicked.connect(lambda: self._find(backward=True))
        self.btn_next.clicked.connect(lambda: self._find(backward=False))
        self.btn_rep.clicked.connect(self._replace_once)
        self.btn_all.clicked.connect(self._replace_all)

    def _flags(self):
        flags = QtGui.QTextDocument.FindFlag(0)
        if self.ck_case.isChecked(): flags |= QtGui.QTextDocument.FindFlag.FindCaseSensitively
        if self.ck_word.isChecked(): flags |= QtGui.QTextDocument.FindFlag.FindWholeWords
        return flags

    def _find(self, backward=False):
        needle = self.find_edit.text()
        if not needle: return
        flags = self._flags()
        if backward: flags |= QtGui.QTextDocument.FindFlag.FindBackward
        found = self.editor.find(needle, flags)
        if not found:
            c = self.editor.textCursor()
            c.movePosition(QtGui.QTextCursor.MoveOperation.End if backward else QtGui.QTextCursor.MoveOperation.Start)
            self.editor.setTextCursor(c)
            self.editor.find(needle, flags)

    def _replace_once(self):
        needle = self.find_edit.text()
        rep = self.replace_edit.text()
        if not needle: return
        cur = self.editor.textCursor()
        if cur.hasSelection() and cur.selectedText() == needle: cur.insertText(rep)
        self._find(backward=False)

    def _replace_all(self):
        needle = self.find_edit.text()
        rep = self.replace_edit.text()
        if not needle: return
        flags = self._flags()
        doc = self.editor.document()
        c = QtGui.QTextCursor(doc)
        c.beginEditBlock()
        c.movePosition(QtGui.QTextCursor.MoveOperation.Start)
        self.editor.setTextCursor(c)
        count = 0
        while self.editor.find(needle, flags):
            cur = self.editor.textCursor()
            cur.insertText(rep)
            count += 1
        c.endEditBlock()
        QtWidgets.QMessageBox.information(self, "Replace All", f"Replaced {count} occurrence(s).")


class HindiEditor(QtWidgets.QTextEdit):
    countsChanged = QtCore.Signal(int, int)
    cursorPositionChangedDetailed = QtCore.Signal(int, int)
    contextActionTriggered = QtCore.Signal(str) 
    englishModeToggled = QtCore.Signal(bool)
    
    suggestionsReady = QtCore.Signal(list)
    navigateSuggestion = QtCore.Signal(int)
    insertSuggestionTrigger = QtCore.Signal()

    def __init__(self, translit: AdaptiveTransliterator, state: AppState, parent=None):
        super().__init__(parent)
        self.translit = translit
        self.state = state
        self.setAcceptRichText(True)
        f = QtGui.QFont(self.state.font_family, self.state.font_size)
        self.setFont(f)
        self.document().setDefaultFont(f)
        self.setPlaceholderText("Type phonetically (Latin)... \nDevanagari will appear.")

        self._composing_latin = ""
        self._composing_start_pos: int | None = None
        self._composing_display_len = 0
        self._ignore_cursor_move = False
        self._english_mode = False
        self._has_dock_suggestions = False
        self._last_word = "" 
        
        self.fuzzy_thread = None

        self.count_timer = QtCore.QTimer(self)
        self.count_timer.setSingleShot(True)
        self.count_timer.setInterval(500)
        self.count_timer.timeout.connect(self._emit_counts_actual)
        
        self.sugg_popup = QtWidgets.QListWidget()
        self.sugg_popup.setWindowFlags(Qt.WindowType.ToolTip | Qt.WindowType.FramelessWindowHint)
        self.sugg_popup.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground) 
        self.sugg_popup.setFocusPolicy(Qt.FocusPolicy.NoFocus)
        self.sugg_popup.setObjectName("suggPopup")
        self.sugg_popup.setFixedWidth(240) 
        self.sugg_popup.itemPressed.connect(self._insert_selected_suggestion)
        
        self.textChanged.connect(self._on_text_changed)
        self.cursorPositionChanged.connect(self._on_cursor_moved)

    def focusOutEvent(self, e):
        self.sugg_popup.hide()
        super().focusOutEvent(e)

    def set_has_dock_suggestions(self, val: bool):
        self._has_dock_suggestions = val

    def _get_word_before_cursor(self) -> str:
        cur = self.textCursor()
        if cur.position() > 0 and self.toPlainText()[cur.position() - 1].isspace():
            cur.movePosition(QtGui.QTextCursor.MoveOperation.Left)
        cur.movePosition(QtGui.QTextCursor.MoveOperation.StartOfWord, QtGui.QTextCursor.MoveMode.KeepAnchor)
        return cur.selectedText().strip()

    def _process_suggestions(self, prefix: str):
        if not prefix or self._english_mode:
            self.sugg_popup.hide()
            self.suggestionsReady.emit([])
            return
            
        prefix_lower = prefix.lower()
        search_prefix = prefix_lower.rstrip('्') if prefix_lower.endswith('्') else prefix_lower
        all_words = list(self.state.suggestion_words)
        
        matches = [w for w in all_words if w.lower().startswith(search_prefix)]
        
        predicted_words = set()
        if self._last_word:
            last_lower = self._last_word.lower()
            if last_lower in self.state.next_words:
                predicted_words = set(self.state.next_words[last_lower].keys())
        
        matches.sort(key=lambda w: (
            w.lower() not in predicted_words, 
            w.lower() != prefix_lower,        
            w.lower() != search_prefix,       
            len(w),                           
            w.lower()                         
        ))
        
        if matches:
            if self.state.suggestion_mode in ["Inline", "Both"]:
                self._process_suggestions_direct(matches)
            else:
                self.sugg_popup.hide()
            if self.state.suggestion_mode in ["Dock", "Both"]:
                self.suggestionsReady.emit(matches[:100])
            else:
                self.suggestionsReady.emit([])
                
        elif len(prefix_lower) >= 3:
            if self.fuzzy_thread and self.fuzzy_thread.isRunning():
                self.fuzzy_thread.terminate()
            self.fuzzy_thread = FuzzySearchThread(prefix_lower, all_words, parent=self)
            self.fuzzy_thread.result_ready.connect(self._on_fuzzy_results)
            self.fuzzy_thread.start()
        else:
            self.sugg_popup.hide()
            self.suggestionsReady.emit([])

    @QtCore.Slot(list)
    def _on_fuzzy_results(self, matches: list):
        if self.state.suggestion_mode in ["Inline", "Both"] and matches:
            self._process_suggestions_direct(matches)
        else:
            self.sugg_popup.hide()

        if self.state.suggestion_mode in ["Dock", "Both"] and matches:
            self.suggestionsReady.emit(matches[:100])

    def _process_suggestions_direct(self, words: list):
        if not words: return
        self.sugg_popup.clear()
        
        main_win = self.window()
        columns = getattr(main_win, 'sugg_columns', 1)
        rows = getattr(main_win, 'sugg_rows', 10)
        show_nums = getattr(main_win, 'sugg_show_numbers', True)
        max_words = columns * rows

        self.sugg_popup.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.sugg_popup.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        
        if columns > 1:
            self.sugg_popup.setFlow(QtWidgets.QListView.Flow.TopToBottom)
            self.sugg_popup.setWrapping(True)
            self.sugg_popup.setResizeMode(QtWidgets.QListView.ResizeMode.Adjust)
        else:
            self.sugg_popup.setFlow(QtWidgets.QListView.Flow.TopToBottom)
            self.sugg_popup.setWrapping(False)

        for i, word in enumerate(words[:max_words]):
            item = QtWidgets.QListWidgetItem()
            if show_nums and i < 11:
                if i == 0: display = f"[Tab] {word}" 
                elif i < 10: display = f"[{i}] {word}"
                else: display = f"[0] {word}"
            else:
                display = word
                
            item.setText(display)
            item.setData(Qt.ItemDataRole.UserRole, word)
            self.sugg_popup.addItem(item)
            
        self.sugg_popup.setCurrentRow(0)
        self._calculate_and_show_popup(min(len(words), max_words), rows, columns)

    def _calculate_and_show_popup(self, actual_items, rows, columns):
        rect = self.cursorRect()
        pt = self.viewport().mapToGlobal(rect.bottomLeft())
        pt.setY(pt.y() + 5)
        
        item_h = self.sugg_popup.sizeHintForRow(0)
        if item_h <= 0: item_h = 32
        actual_rows = min(rows, actual_items)
        frame_offset = self.sugg_popup.frameWidth() * 2
        total_h = (item_h * actual_rows) + frame_offset + 6
        
        col_width = 170
        if columns > 1:
            self.sugg_popup.setGridSize(QtCore.QSize(col_width, item_h))
            actual_cols = math.ceil(actual_items / rows)
            actual_cols = max(1, min(actual_cols, columns))
            total_w = (col_width * actual_cols) + frame_offset + 6
            self.sugg_popup.setFixedWidth(total_w)
        else:
            self.sugg_popup.setGridSize(QtCore.QSize())
            self.sugg_popup.setFixedWidth(240)

        self.sugg_popup.setFixedHeight(total_h)
        
        screen = QtGui.QGuiApplication.screenAt(pt)
        if screen:
            screen_geom = screen.availableGeometry()
            if pt.x() + self.sugg_popup.width() > screen_geom.right():
                pt.setX(screen_geom.right() - self.sugg_popup.width() - 10)
            if pt.y() + self.sugg_popup.height() > screen_geom.bottom():
                pt.setY(self.viewport().mapToGlobal(rect.topLeft()).y() - self.sugg_popup.height() - 5)

        self.sugg_popup.move(pt)
        self.sugg_popup.show()

    def keyPressEvent(self, ev: QtGui.QKeyEvent):
        key, mods = ev.key(), ev.modifiers()
        
        if key == Qt.Key.Key_Space and (mods & Qt.KeyboardModifier.ControlModifier):
            self._english_mode = not self._english_mode
            if self._composing_latin: 
                self._commit_composing()
            self.sugg_popup.hide()
            self.englishModeToggled.emit(self._english_mode)
            QtWidgets.QToolTip.showText(self.mapToGlobal(self.cursorRect().bottomRight()),
                                        "English (Hinglish) Mode ON" if self._english_mode else "Hindi Translit Mode ON",
                                        self)
            return

        if self.sugg_popup.isVisible():
            if key == Qt.Key.Key_Up:
                r = self.sugg_popup.currentRow()
                if r > 0: self.sugg_popup.setCurrentRow(r - 1)
                return
            elif key == Qt.Key.Key_Down:
                r = self.sugg_popup.currentRow()
                if r < self.sugg_popup.count() - 1: self.sugg_popup.setCurrentRow(r + 1)
                return
            elif key == Qt.Key.Key_Tab: 
                item = self.sugg_popup.currentItem()
                if item:
                    self._insert_suggestion_text(item.data(Qt.ItemDataRole.UserRole))
                return
            elif key in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                self.sugg_popup.hide()
            elif key == Qt.Key.Key_Escape:
                self.sugg_popup.hide()
                return

        if mods & Qt.KeyboardModifier.ControlModifier: 
            if self.state.suggestion_mode in ["Dock", "Both"] and self._has_dock_suggestions:
                if key == Qt.Key.Key_Up or key == Qt.Key.Key_Down:
                    step = -1 if key == Qt.Key.Key_Up else 1
                    self.navigateSuggestion.emit(step)
                    return

            if self.sugg_popup.isVisible() and Qt.Key.Key_0 <= key <= Qt.Key.Key_9:
                num = key - Qt.Key.Key_0
                idx = 10 if num == 0 else num
                if idx < self.sugg_popup.count():
                    item = self.sugg_popup.item(idx)
                    if item:
                        self._insert_suggestion_text(item.data(Qt.ItemDataRole.UserRole))
                return
            
            return super().keyPressEvent(ev)

        if self.state.suggestion_mode in ["Dock", "Both"] and self._has_dock_suggestions and not self.sugg_popup.isVisible():
            if key == Qt.Key.Key_Tab:
                self.insertSuggestionTrigger.emit()
                return

        if self._english_mode:
            self.sugg_popup.hide()
            return super().keyPressEvent(ev)
            
        elif key == Qt.Key.Key_Space:
            current_word = ""
            
            if self._composing_latin: 
                current_word = self.translit.translit_token(self._composing_latin) if self.state.transliteration_enabled else self._composing_latin
                self._commit_composing()
            else:
                current_word = self._get_word_before_cursor()
                
            if current_word:
                self.state.learn_bigram(self._last_word, current_word)
                self._last_word = current_word

            super().keyPressEvent(ev)
            self._show_next_word_predictions()
            return

        if key in {Qt.Key.Key_Left, Qt.Key.Key_Right, Qt.Key.Key_Up, Qt.Key.Key_Down, Qt.Key.Key_Home, Qt.Key.Key_End}:
            if self._composing_latin: self._commit_composing()
            return super().keyPressEvent(ev)

        if key == Qt.Key.Key_Backspace:
            if self._composing_latin:
                self._composing_latin = self._composing_latin[:-1]
                if self._composing_latin: 
                    self._update_composing()
                    self._process_suggestions(self.translit.translit_token(self._composing_latin) if self.state.transliteration_enabled else self._composing_latin)
                else: 
                    self._remove_composing()
                    self.sugg_popup.hide()
                    self.suggestionsReady.emit([])
                return
            return super().keyPressEvent(ev)

        if key in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
            if self._composing_latin: self._commit_composing()
            return super().keyPressEvent(ev)

        txt = ev.text()
        if not txt or re.search(r"[\u0900-\u097F]", txt):
            if self._composing_latin: self._commit_composing()
            return super().keyPressEvent(ev)

        if txt.isprintable() and len(txt) == 1 and not txt.isspace() and ord(txt) < 128:
            cur = self.textCursor()
            if not self._composing_latin:
                if cur.hasSelection(): cur.removeSelectedText()
                start = cur.position()
                self._composing_start_pos = start
                self._composing_latin = txt
                
                translit_text = self.translit.translit_token(self._composing_latin) if self.state.transliteration_enabled else self._composing_latin
                fmt = cur.charFormat()
                
                self._ignore_cursor_move = True
                cur.insertText(translit_text, fmt)
                self._composing_display_len = len(translit_text)
                self.setTextCursor(cur)
                self._ignore_cursor_move = False
                
                self._process_suggestions(translit_text)
                return
            else:
                self._composing_latin += txt
                translit_text = self.translit.translit_token(self._composing_latin) if self.state.transliteration_enabled else self._composing_latin
                self._update_composing_with_text(translit_text)
                
                self._process_suggestions(translit_text)
                return

        if self._composing_latin: self._commit_composing()
        return super().keyPressEvent(ev)

    def _show_next_word_predictions(self):
        if not self._english_mode and self._last_word:
            last_lower = self._last_word.lower()
            if last_lower in self.state.next_words:
                predictions = sorted(self.state.next_words[last_lower].items(), key=lambda x: x[1], reverse=True)
                top_words = [p[0] for p in predictions][:20]
                if top_words:
                    self._process_suggestions_direct(top_words)
                    return
        self.sugg_popup.hide()

    def _on_text_changed(self): 
        self.count_timer.start()

    def _emit_counts_actual(self):
        txt = self.toPlainText()
        w = len([w for w in re.split(r"\s+", txt.strip()) if w]) if txt.strip() else 0
        self.countsChanged.emit(w, len(txt))

    def _on_cursor_moved(self):
        if self._ignore_cursor_move: return
        if self._composing_latin: self._commit_composing()
        
        if not self._english_mode:
            cur = self.textCursor()
            cur.movePosition(QtGui.QTextCursor.MoveOperation.StartOfWord, QtGui.QTextCursor.MoveMode.KeepAnchor)
            self._process_suggestions(cur.selectedText().strip())
        
        cur = self.textCursor()
        line = cur.blockNumber() + 1
        col = cur.positionInBlock()
        self.cursorPositionChangedDetailed.emit(line, col)

    def _insert_suggestion_text(self, word: str):
        cur = self.textCursor()
        self._ignore_cursor_move = True
        
        if self._composing_latin: 
            self._remove_composing()
        else:
            cur.movePosition(QtGui.QTextCursor.MoveOperation.StartOfWord, QtGui.QTextCursor.MoveMode.KeepAnchor)
            cur.removeSelectedText()
            
        cur.insertText(word)
        self.setTextCursor(cur)
        self.setFocus()
        
        if self._last_word:
            self.state.learn_bigram(self._last_word, word)
        self._last_word = word
        
        self.suggestionsReady.emit([])
        self._composing_latin = ""
        self._ignore_cursor_move = False
        
        super().insertPlainText(" ")
        self._show_next_word_predictions()

    def insert_suggestion_from_dock(self, word: str):
        self._insert_suggestion_text(word)

    def _insert_selected_suggestion(self, item):
        if item:
            self._insert_suggestion_text(item.data(Qt.ItemDataRole.UserRole))

    def _update_composing_with_text(self, translit_text: str):
        if self._composing_start_pos is None: return
        self._ignore_cursor_move = True
        cur = self.textCursor()
        cur.setPosition(self._composing_start_pos)
        cur.setPosition(self._composing_start_pos + self._composing_display_len, QtGui.QTextCursor.MoveMode.KeepAnchor)
        fmt = cur.charFormat()
        cur.insertText(translit_text, fmt)
        self._composing_display_len = len(translit_text)
        self.setTextCursor(cur)
        self._ignore_cursor_move = False

    def _update_composing(self):
        t_text = self.translit.translit_token(self._composing_latin) if self.state.transliteration_enabled else self._composing_latin
        self._update_composing_with_text(t_text)

    def _remove_composing(self):
        if self._composing_start_pos is None: return
        self._ignore_cursor_move = True
        cur = self.textCursor()
        cur.setPosition(self._composing_start_pos)
        cur.setPosition(self._composing_start_pos + self._composing_display_len, QtGui.QTextCursor.MoveMode.KeepAnchor)
        cur.removeSelectedText()
        self._composing_latin = ""
        self._composing_start_pos = None
        self._composing_display_len = 0
        self.setTextCursor(cur)
        self._ignore_cursor_move = False

    def _commit_composing(self):
        if not self._composing_latin: return
        self._composing_latin = ""
        self._composing_start_pos = None
        self._composing_display_len = 0
        self.count_timer.start()

    def paste(self):
        md = QtWidgets.QApplication.clipboard().mimeData()
        text = md.text() if md else ""
        if not text: return super().paste()
        if re.search(r"[\u0900-\u097F]", text):
            if self._composing_latin: self._commit_composing()
            return super().paste()
            
        if self._english_mode:
            return super().paste()
            
        parts = re.split(r"(\s+)", text)
        for p in parts:
            if p.strip() == "": super().insertPlainText(p)
            else:
                chunk = self.translit.translit_token(p) if self.state.transliteration_enabled else p
                super().insertPlainText(chunk)
        self.count_timer.start()


    def contextMenuEvent(self, e):
        menu = self.createStandardContextMenu()
        
        menu.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, False)
        
        main_window = self.window()
        is_dark = getattr(main_window, 'dark_mode_enabled', False)
        
        if is_dark:
            menu.setStyleSheet("""
                QMenu { background-color: #1f2937; color: #f9fafb; border: 1px solid #4b5563; padding: 4px; }
                QMenu::item:selected { background-color: #059669; }
                QMenu::separator { background-color: #4b5563; height: 1px; margin: 4px 0px; }
            """)
        else:
            menu.setStyleSheet("""
                QMenu { background-color: #ffffff; color: #111827; border: 1px solid #d1d5db; padding: 4px; }
                QMenu::item:selected { background-color: #059669; color: white; }
                QMenu::separator { background-color: #e5e7eb; height: 1px; margin: 4px 0px; }
            """)

        menu.addSeparator()
        
        cur = self.textCursor()
        
        # IMAGE CONTEXT MENU
        char_fmt = cur.charFormat()
        if char_fmt.isImageFormat():
            img_menu = menu.addMenu("Image Formatting")
            img_menu.addAction("Resize Image...", lambda: self.contextActionTriggered.emit("resize_image"))
            menu.addSeparator()

        # TABLE CONTEXT MENU
        table = cur.currentTable()
        if table:
            t_menu = menu.addMenu("Table Formatting")
            t_menu.addAction("Add Row Below", lambda: table.insertRows(table.cellAt(cur).row() + 1, 1))
            t_menu.addAction("Add Row Above", lambda: table.insertRows(table.cellAt(cur).row(), 1))
            t_menu.addAction("Add Column Right", lambda: table.insertColumns(table.cellAt(cur).column() + 1, 1))
            t_menu.addAction("Merge Selected Cells", lambda: table.mergeCells(cur))
            t_menu.addAction("Split/Unmerge Cell", lambda: table.splitCell(table.cellAt(cur).row(), table.cellAt(cur).column(), 1, 1))
            t_menu.addSeparator()
            t_menu.addAction("Toggle Borders (On/Off)", lambda: self.contextActionTriggered.emit("toggle_borders"))
            
            # Structural Deletion Actions
            t_menu.addSeparator()
            t_menu.addAction("Delete Current Row", lambda: table.removeRows(table.cellAt(cur).row(), 1))
            t_menu.addAction("Delete Current Column", lambda: table.removeColumns(table.cellAt(cur).column(), 1))
            
            menu.addSeparator()

        if table:
            t_menu = menu.addMenu("Table Alignment (Page)")
            t_menu.addAction("Align Left", lambda: self.contextActionTriggered.emit("table_left"))
            t_menu.addAction("Align Center", lambda: self.contextActionTriggered.emit("table_center"))
            t_menu.addAction("Align Right", lambda: self.contextActionTriggered.emit("table_right"))
            menu.addSeparator()

        menu.addAction("Train map from selection", lambda: self.contextActionTriggered.emit("train_sel"))
        menu.addAction("Edit Correction Dictionary", lambda: self.contextActionTriggered.emit("edit_dict"))
        menu.addSeparator()
        menu.addAction("New Document", lambda: self.contextActionTriggered.emit("new_doc"))
        menu.addAction("Open...", lambda: self.contextActionTriggered.emit("open_doc"))
        menu.addAction("Export as PDF", lambda: self.contextActionTriggered.emit("export_pdf"))
        
        sel = cur.selectedText()
        if sel:
            matches = [(k, v) for k, v in self.state.user_dict.items() if k in sel]
            if matches:
                menu.addSeparator()
                sub = menu.addMenu("Apply correction")
                
                # Limit to the first 10 matches to prevent menu bloat
                for k, v in matches[:10]:
                    a = sub.addAction(f"{k} → {v}")
                    a.triggered.connect(lambda checked=False, kk=k, vv=v: self._apply_correction_in_selection(kk, vv))
                
                # Add a visual indicator if there are hidden matches
                if len(matches) > 10:
                    sub.addAction(f"... and {len(matches) - 10} more").setEnabled(False)
                    
                sub.addSeparator()
                
                # Option to convert all matches at once
                a_all = sub.addAction("Convert All Matches")
                a_all.triggered.connect(lambda checked=False: self.contextActionTriggered.emit("context_correct_all"))
                
            menu.addSeparator()
            translate_action = menu.addAction("Translate Selection to Hindi")
            translate_action.triggered.connect(lambda checked=False: self.contextActionTriggered.emit("context_translate"))


        # Render menu
        menu.exec(e.globalPos())

    def _apply_correction_in_selection(self, key, val):
        cur = self.textCursor()
        if not cur.hasSelection(): cur.select(QtGui.QTextCursor.SelectionType.WordUnderCursor)
        text = cur.selectedText()
        cur.insertText(text.replace(key, val))

    def _align_table(self, alignment):
        cur = self.textCursor()
        table = cur.currentTable()
        if table:
            fmt = table.format()
            fmt.setAlignment(alignment)
            table.setFormat(fmt)
            self.textChanged.emit()


class MainWindow(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle(APP_NAME)
        self.resize(1200, 820)
        self.setWindowIcon(QIcon("icons/icon.png"))        
        self.state = AppState()
        self.state.load()
        
        self.current_filepath: str | None = None
        
        self.settings = QtCore.QSettings("TranslitStudio", "TranslitDocMaker")
        self.dark_mode_enabled = self.settings.value("dark_mode", True, type=bool)
        
        self.sugg_style = self.settings.value("sugg_style", "Google (Search Style)", type=str)
        self.sugg_font_size = int(self.settings.value("sugg_font_size", 14))
        self.sugg_show_numbers = self.settings.value("sugg_show_numbers", True, type=bool)
        
        self.sugg_columns = int(self.settings.value("sugg_columns", 1))
        self.sugg_rows = int(self.settings.value("sugg_rows", 10))
        self.sugg_spacing = self.settings.value("sugg_spacing", "Normal", type=str)
        
        self.sugg_text_color = self.settings.value("sugg_text_color", "", type=str)
        self.sugg_bold = self.settings.value("sugg_bold", False, type=bool)
        
        self.printer = QPrinter(QPrinter.PrinterMode.HighResolution)
        
        self.autosave_enabled: bool = False
        self.autosave_timer = QTimer(self)
        self.autosave_timer.setInterval(60000)
        self.autosave_timer.timeout.connect(self._perform_autosave)
        self.is_dirty = False

        self.app_seconds = 0
        self.doc_seconds = 0
        self.doc_timer_active = False
        self.global_timer = QTimer(self)
        self.global_timer.timeout.connect(self._on_timer_tick)
        self.global_timer.start(1000)

        try: self.translit = AdaptiveTransliterator(self.state)
        except RuntimeError as e:
            QtWidgets.QMessageBox.critical(None, "Missing dependency", str(e))
            raise SystemExit(1)

        self.workspace_scroll = QtWidgets.QScrollArea()
        self.workspace_scroll.setAlignment(Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignTop)
        self.workspace_scroll.setWidgetResizable(True)
        
        self.canvas_container = QtWidgets.QWidget()
        self.canvas_container.setStyleSheet("background-color: transparent;")
        self.canvas_layout = QtWidgets.QVBoxLayout(self.canvas_container)
        self.canvas_layout.setAlignment(Qt.AlignmentFlag.AlignHCenter | Qt.AlignmentFlag.AlignTop)
        self.canvas_layout.setContentsMargins(0, 0, 0, 0)
        
        self.page_frame = QtWidgets.QFrame()
        self.page_frame.setObjectName("page_frame")
        self.page_layout = QtWidgets.QVBoxLayout(self.page_frame)
        self.page_layout.setContentsMargins(0, 0, 0, 0)
        self.page_layout.setSpacing(0)

        self.editor = HindiEditor(self.translit, self.state, self)
        
        # Safe localized event filter
        self.editor.installEventFilter(self)
        
        self.page_layout.addWidget(self.editor)
        self.canvas_layout.addWidget(self.page_frame)
        
        self.workspace_scroll.setWidget(self.canvas_container)
        self.setCentralWidget(self.workspace_scroll)
        
        self.workspace_scroll.verticalScrollBar().valueChanged.connect(lambda _: self.editor.sugg_popup.hide())
        self.workspace_scroll.horizontalScrollBar().valueChanged.connect(lambda _: self.editor.sugg_popup.hide())


        self.editor.cursorPositionChanged.connect(self._ensure_cursor_visible)

        self.editor.countsChanged.connect(self._update_counts)
        self.editor.cursorPositionChangedDetailed.connect(self._update_cursor_position_label)
        self.editor.contextActionTriggered.connect(self._handle_context_action)
        self.editor.textChanged.connect(self._mark_dirty)
        self.editor.englishModeToggled.connect(self._update_translit_badge)

        self.finddlg = FindReplaceDialog(self.editor, self)
        
        self.sr_worker: SpeechWorker | None = None
        self.sr_listening = False
        self._ctrl_l_down = False
        
        self.export_thread = None
        
        self._setup_actions() 
        self._build_toolbars()
        self._build_menus()
        self._build_sidebar()
        self._build_suggestions_sidebar()
        
        self.editor.suggestionsReady.connect(self._populate_dock_suggestions)
        self.editor.navigateSuggestion.connect(self._navigate_dock_suggestion_list)
        self.editor.insertSuggestionTrigger.connect(self._insert_selected_dock_suggestion)
        self.editor.currentCharFormatChanged.connect(self._lock_font_properties)

        self._build_statusbar()
        self._apply_theme()
        self._update_translit_badge()
        self._update_view_mode()
        
        QTimer.singleShot(0, lambda: self._new_file(prompt_autosave=False, is_startup=True))

    def _lock_font_properties(self, fmt):
        # Prevents the editor from resetting to default system fonts when deleting text
        if hasattr(self, 'state') and (self.editor.fontFamily() != self.state.font_family or self.editor.fontPointSize() != self.state.font_size):
            self.editor.blockSignals(True)
            self.editor.setFontFamily(self.state.font_family)
            self.editor.setFontPointSize(self.state.font_size)
            self.editor.blockSignals(False)

    def _translate_selection(self):
        cur = self.editor.textCursor()
        if not cur.hasSelection():
            # If nothing selected, select the current word
            cur.select(QtGui.QTextCursor.SelectionType.WordUnderCursor)
            
        text = cur.selectedText().strip()
        if not text: return
        
        self.status.showMessage("Translating...")
        QtWidgets.QApplication.processEvents()
        
        try:
            # Ping Google Translate API
            url = "https://translate.googleapis.com/translate_a/single?client=gtx&sl=en&tl=hi&dt=t&q=" + urllib.parse.quote(text)
            req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
            response = urllib.request.urlopen(req, timeout=3)
            data = json.loads(response.read().decode('utf-8'))
            translated_text = "".join([sentence[0] for sentence in data[0]])
            
            cur.insertText(translated_text)
            self.status.showMessage("Translated successfully.", 2000)
        except Exception as e:
            self.status.showMessage("Translation failed. Check internet connection.", 3000)
            logging.error(f"Translation error: {e}")
        
    def _modify_indent(self, delta):
        cur = self.editor.textCursor()
        fmt = cur.blockFormat()
        new_indent = max(0, fmt.indent() + delta)
        fmt.setIndent(new_indent)
        cur.setBlockFormat(fmt)
        
    def _set_line_spacing(self, percentage):
        cur = self.editor.textCursor()
        fmt = cur.blockFormat()
        
        # 1 is the strict PySide6 integer enum for ProportionalHeight.
        # This guarantees 50 = 50% spacing, 100 = 100% spacing, etc.
        fmt.setLineHeight(float(percentage), 1) 
        
        cur.setBlockFormat(fmt)
        self.editor.setFocus()     

    def _insert_custom_numbered(self, list_style, prefix="", suffix="."):
            fmt = QtGui.QTextListFormat()
            fmt.setStyle(list_style)
            fmt.setNumberPrefix(prefix)
            fmt.setNumberSuffix(suffix)
            self.editor.textCursor().createList(fmt)        
        

    def _change_sugg_columns(self):
        val, ok = QtWidgets.QInputDialog.getInt(self, "Popup Layout", "Number of Columns:", self.sugg_columns, 1, 5)
        if ok:
            self.sugg_columns = val
            self.settings.setValue("sugg_columns", val)
            self._apply_popup_style()

    def _change_sugg_rows(self):
        val, ok = QtWidgets.QInputDialog.getInt(self, "Popup Layout", "Words per Column (Rows):", self.sugg_rows, 3, 30)
        if ok:
            self.sugg_rows = val
            self.settings.setValue("sugg_rows", val)
            self._apply_popup_style()

    def _change_sugg_spacing(self, spacing_level):
        self.sugg_spacing = spacing_level
        self.settings.setValue("sugg_spacing", spacing_level)
        self._apply_popup_style()

    def _change_sugg_max_words(self):
        val, ok = QtWidgets.QInputDialog.getInt(self, "Max Words", "Number of suggestions to show:", self.sugg_max_words, 3, 30)
        if ok:
            self.sugg_max_words = val
            self.settings.setValue("sugg_max_words", val)

    def _change_sugg_color(self):
        col = QtWidgets.QColorDialog.getColor(parent=self, title="Choose Suggestion Text Color")
        if col.isValid():
            self.sugg_text_color = col.name()
            self.settings.setValue("sugg_text_color", self.sugg_text_color)
            self._apply_popup_style()

    def _reset_sugg_color(self):
        self.sugg_text_color = ""
        self.settings.setValue("sugg_text_color", "")
        self._apply_popup_style()

    def _toggle_sugg_bold(self, checked):
        self.sugg_bold = checked
        self.settings.setValue("sugg_bold", checked)
        self._apply_popup_style()

    def _toggle_sugg_numbers(self, checked):
        self.sugg_show_numbers = checked
        self.settings.setValue("sugg_show_numbers", checked)

    def _change_popup_style(self, style_name):
        self.sugg_style = style_name
        self.settings.setValue("sugg_style", style_name)
        self._apply_popup_style()

    def _change_popup_font_size(self):
        size, ok = QtWidgets.QInputDialog.getInt(self, "Popup Font Size", "Enter font size (px):", self.sugg_font_size, 8, 48)
        if ok:
            self.sugg_font_size = size
            self.settings.setValue("sugg_font_size", size)
            self._apply_popup_style()

    def _apply_popup_style(self):
        if self.sugg_style == "Default (OS Native)":
            self.editor.sugg_popup.setStyleSheet("")
            font = self.editor.sugg_popup.font()
            font.setPointSize(self.sugg_font_size)
            font.setBold(self.sugg_bold)
            self.editor.sugg_popup.setFont(font)
            return

        css = style.get_popup_style(
            self.sugg_style, 
            self.dark_mode_enabled, 
            self.sugg_text_color, 
            self.sugg_font_size, 
            self.sugg_bold, 
            self.sugg_spacing
        )
        self.editor.sugg_popup.setStyleSheet(css)

    def _on_timer_tick(self):
        self.app_seconds += 1
        if self.doc_timer_active:
            self.doc_seconds += 1

        app_m, app_s = divmod(self.app_seconds, 60)
        app_h, app_m = divmod(app_m, 60)
        self.lbl_app_time.setText(f"App: {app_h:02d}:{app_m:02d}:{app_s:02d}")

        doc_m, doc_s = divmod(self.doc_seconds, 60)
        doc_h, doc_m = divmod(doc_m, 60)
        self.lbl_doc_time.setText(f"Doc: {doc_h:02d}:{doc_m:02d}:{doc_s:02d}")

        if hasattr(self, 'clock_label'):
            self.clock_label.setText(datetime.now().strftime("%H:%M"))

    def resizeEvent(self, event):
        self.editor.sugg_popup.hide()
        super().resizeEvent(event)

    def moveEvent(self, event):
        self.editor.sugg_popup.hide()
        super().moveEvent(event)

    def eventFilter(self, watched, event):
        try:
            if event.type() == QtCore.QEvent.Type.KeyPress:
                key_event: QtGui.QKeyEvent = event
                if key_event.key() == Qt.Key.Key_F11:
                    if self.isFullScreen(): self.showNormal()
                    else: self.showFullScreen()
                    return True
                    
                if key_event.key() == Qt.Key.Key_L and (key_event.modifiers() & Qt.KeyboardModifier.ControlModifier) and not key_event.isAutoRepeat():
                    if not self.sr_listening:
                        self._start_sr_listening()
                    self._ctrl_l_down = True
                    return True
            elif event.type() == QtCore.QEvent.Type.KeyRelease:
                key_event: QtGui.QKeyEvent = event
                if key_event.key() == Qt.Key.Key_L and not key_event.isAutoRepeat():
                    if self.sr_listening and self._ctrl_l_down:
                        self._stop_sr_listening_and_insert()
                    self._ctrl_l_down = False
                    return True
        except Exception:
            pass
        return super().eventFilter(watched, event)

    def _start_sr_listening(self):
        if not HAS_SR:
            QtWidgets.QMessageBox.information(self, "Speech missing", "Install SpeechRecognition and PyAudio")
            return
        if self.sr_worker and self.sr_worker.isRunning(): return
        self.sr_worker = SpeechWorker(lang="hi-IN", parent=self)
        self.sr_worker.partial.connect(self._update_speech_label_partial)
        self.sr_worker.finished_text.connect(self._sr_finished)
        self.sr_worker.error.connect(lambda m: self.status.showMessage(f"Speech error: {m}", 4000))
        self.sr_worker.start()
        self.sr_listening = True
        self.lbl_speech_indicator.setText("Listening... release Ctrl+L to stop")

    @QtCore.Slot(str)
    def _update_speech_label_partial(self, s: str):
        if s: self.lbl_speech_indicator.setText(f"Listening... '{s[:30]}...'")
        else: self.lbl_speech_indicator.setText("Listening...")

    def _stop_sr_listening_and_insert(self):
        try:
            if self.sr_worker:
                self.sr_worker.stop()
                self.sr_worker.wait(3000)
        except Exception: pass
        self.sr_listening = False
        self.lbl_speech_indicator.setText("Processing...")

    @QtCore.Slot(str)
    def _sr_finished(self, text: str):
        if text:
            try: devanagari = self.translit.translit_full(text)
            except Exception: devanagari = text
            cur = self.editor.textCursor()
            cur.insertText(devanagari)
            self.status.showMessage("Inserted speech", 3000)
        else:
            self.status.showMessage("No speech recognized", 3000)
        self.sr_worker = None
        self.lbl_speech_indicator.setText("Hold Ctrl+L to speak")

    def _mark_dirty(self):
        self.is_dirty = True
        if not self.doc_timer_active:
            self.doc_timer_active = True

    def _perform_autosave(self):
        if self.autosave_enabled and self.current_filepath and self.is_dirty:
            try:
                ext = self.current_filepath.lower().split('.')[-1]
                if ext == "docx" and HAS_DOCX: 
                    pass 
                elif ext in ["html", "htm"]:
                    self._save_html_silent(self.current_filepath)
                else:
                    save_json_atomically({"content": self.editor.toPlainText()}, Path(self.current_filepath))
                        
                self.state.save_next_words()
                self.is_dirty = False
                self.status.showMessage(f"Autosaved to {os.path.basename(self.current_filepath)}", 2000)
            except Exception: pass

    def _handle_context_action(self, action: str):
        if action == "train_sel": self._train_manual()
        elif action == "edit_dict": self._open_corrections()
        elif action == "new_doc": self.act_new.trigger()
        elif action == "open_doc": self.act_open.trigger()
        elif action == "save_docx": self.act_save_docx.trigger()
        elif action == "export_pdf": self.act_export_pdf.trigger()
        elif action == "table_left": self.editor._align_table(Qt.AlignmentFlag.AlignLeft)
        elif action == "table_center": self.editor._align_table(Qt.AlignmentFlag.AlignHCenter)
        elif action == "table_right": self.editor._align_table(Qt.AlignmentFlag.AlignRight)
        elif action == "toggle_borders":
            table = self.editor.textCursor().currentTable()
            if table:
                fmt = table.format()
                fmt.setBorder(0 if fmt.border() > 0 else 1)
                table.setFormat(fmt)
        elif action == "resize_image":
            cur = self.editor.textCursor()
            fmt = cur.charFormat().toImageFormat()
            if fmt.isValid():
                # Extract clean local path safely handling PySide6 URI prefixes
                img_path = QtCore.QUrl(fmt.name()).toLocalFile() if fmt.name().startswith("file://") else fmt.name()
                
                # Open up a dedicated modification window
                dialog = ImagePropertiesDialog(img_path, fmt.width(), fmt.height(), self)
                if dialog.exec() == QtWidgets.QDialog.DialogCode.Accepted:
                    new_w, new_h, cropped_path = dialog.get_results()
                    
                    # Force editor selection onto the image char to ensure format sticks
                    cur.movePosition(QtGui.QTextCursor.MoveOperation.Left, QtGui.QTextCursor.MoveMode.KeepAnchor)
                    
                    # Update target source metadata if a crop operation occurred
                    if cropped_path:
                        fmt.setName(QtCore.QUrl.fromLocalFile(cropped_path).toString())
                        
                    fmt.setWidth(float(new_w))
                    fmt.setHeight(float(new_h))
                    cur.setCharFormat(fmt)
                    
        elif action == "context_correct_all":
            cur = self.editor.textCursor()
            if not cur.hasSelection(): 
                cur.select(QtGui.QTextCursor.SelectionType.WordUnderCursor)
            text = cur.selectedText()
            if not text: return
            for k in sorted(self.state.user_dict.keys(), key=len, reverse=True):
                if k in text:
                    text = text.replace(k, self.state.user_dict[k])
            cur.insertText(text)
                    
        elif action == "context_translate":
            self._translate_selection()
                

    def _setup_actions(self):
                
        self.act_new = QtGui.QAction("New", self, shortcut="Ctrl+N", triggered=lambda: self._new_file(prompt_autosave=True))
        self.act_open = QtGui.QAction("Open", self, shortcut="Ctrl+O", triggered=self._open_file)
        self.act_save = QtGui.QAction("Save", self, shortcut="Ctrl+S", triggered=self._save_manual)
        self.act_save_html = QtGui.QAction("Save as HTML", self, triggered=self._save_html)
        self.act_save_docx = QtGui.QAction("Save as DOCX", self, triggered=self._save_docx)
        self.act_export_pdf = QtGui.QAction("Export as PDF", self, triggered=self._export_pdf)
        self.act_page_setup = QtGui.QAction("Page Setup...", self, triggered=self._page_setup)
        self.act_print = QtGui.QAction("Print (Preview)", self, shortcut="Ctrl+P", triggered=self._print_doc)
        
        self.act_undo = QtGui.QAction("Undo", self, shortcut="Ctrl+Z", triggered=self.editor.undo)
        self.act_redo = QtGui.QAction("Redo", self, shortcut="Ctrl+Y", triggered=self.editor.redo)
        self.act_find = QtGui.QAction("Find & Replace", self, shortcut="Ctrl+F", triggered=self.finddlg.show)

        self.act_dict = QtGui.QAction("Correction Dictionary", self, shortcut="Ctrl+Shift+C", triggered=self._open_corrections)
        self.act_sugg_manager = QtGui.QAction("Suggestion Database Manager...", self, triggered=self._open_suggestion_trainer)
        self.act_translit = QtGui.QAction("Toggle Translit", self, shortcut="Ctrl+T", triggered=self._toggle_transliteration)
        self.act_theme = QtGui.QAction("Theme", self, shortcut="Ctrl+D", triggered=self._toggle_theme)
        self.act_phrase_manager = QtGui.QAction("Phrase Database Manager...", self, triggered=self._open_phrase_manager)
        
        self.act_pred_manager = QtGui.QAction("Prediction Model Manager...", self, triggered=self._open_prediction_manager)
        
        self.view_grp = QtGui.QActionGroup(self)
        self.act_view_web = QtGui.QAction("Web Layout (Default)", self, checkable=True)
        self.act_view_port = QtGui.QAction("Canvas Layout (Portrait)", self, checkable=True)
        self.act_view_land = QtGui.QAction("Canvas Layout (Landscape)", self, checkable=True)
        for a in (self.act_view_web, self.act_view_port, self.act_view_land):
            self.view_grp.addAction(a)
            a.triggered.connect(self._change_view_mode)
            
        if self.state.view_mode == "Portrait": self.act_view_port.setChecked(True)
        elif self.state.view_mode == "Landscape": self.act_view_land.setChecked(True)
        else: self.act_view_web.setChecked(True)

        self.sugg_grp = QtGui.QActionGroup(self)
        self.act_sugg_dock = QtGui.QAction("Dock Sidebar Only", self, checkable=True)
        self.act_sugg_inline = QtGui.QAction("Inline Popup Only", self, checkable=True)
        self.act_sugg_both = QtGui.QAction("Both Dock and Inline", self, checkable=True)
        for a in (self.act_sugg_dock, self.act_sugg_inline, self.act_sugg_both):
            self.sugg_grp.addAction(a)
            a.triggered.connect(self._change_suggestion_mode)
            
        if self.state.suggestion_mode == "Dock": self.act_sugg_dock.setChecked(True)
        elif self.state.suggestion_mode == "Inline": self.act_sugg_inline.setChecked(True)
        else: self.act_sugg_both.setChecked(True)
        
        self.act_insert_date = QtGui.QAction("Insert Date/Time", self, shortcut="Ctrl+Shift+D", triggered=self._insert_datetime)
        self.act_insert_line = QtGui.QAction("Insert Horizontal Line", self, triggered=self._insert_horizontal_line)
        self.act_word_wrap = QtGui.QAction("Toggle Word Wrap", self, triggered=self._toggle_wrap)
        
        self.act_translate = QtGui.QAction("Translate Selection to Hindi", self, shortcut="Ctrl+Shift+E", triggered=self._translate_selection)
        self.addAction(self.act_translate)

    def _change_suggestion_mode(self):
        if self.act_sugg_dock.isChecked(): 
            self.state.suggestion_mode = "Dock"
            self.sugg_dock.show()
        elif self.act_sugg_inline.isChecked(): 
            self.state.suggestion_mode = "Inline"
            self.sugg_dock.hide()
        else: 
            self.state.suggestion_mode = "Both"
            self.sugg_dock.show()

    def _build_toolbars(self):
        tb_file = QtWidgets.QToolBar("File")
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, tb_file)

        tb_file.addAction(self.act_new)
        tb_file.addAction(self.act_open)
        tb_file.addAction(self.act_save)
        tb_file.addAction(self.act_print)
        tb_file.addSeparator()

        self.tpl_btn = QtWidgets.QToolButton()
        self.tpl_btn.setText("Templates")
        self.tpl_btn.setPopupMode(QtWidgets.QToolButton.ToolButtonPopupMode.InstantPopup)
        self.tpl_menu = QtWidgets.QMenu(self)
        self.tpl_btn.setMenu(self.tpl_menu)
        tb_file.addWidget(self.tpl_btn)
        self.tpl_menu.aboutToShow.connect(self._populate_template_menu)
        tb_file.addSeparator()
        
        tb_file.addAction(self.act_undo)
        tb_file.addAction(self.act_redo)
        tb_file.addAction(self.act_theme)

        self.addToolBarBreak() 
        tb_fmt = QtWidgets.QToolBar("Format")
        self.addToolBar(Qt.ToolBarArea.TopToolBarArea, tb_fmt)

        def add_fmt(text, slot, shortcut=None):
            a = QtGui.QAction(text, self)
            if shortcut: a.setShortcut(shortcut)
            a.triggered.connect(slot)
            tb_fmt.addAction(a)

        self.font_combo = QtWidgets.QFontComboBox()
        try: self.font_combo.setCurrentFont(QtGui.QFont(self.state.font_family))
        except Exception: pass
        self.font_combo.currentFontChanged.connect(self._font_changed)
        tb_fmt.addWidget(self.font_combo)

        self.font_size_spin = QtWidgets.QSpinBox()
        self.font_size_spin.setRange(8, 72); self.font_size_spin.setValue(self.state.font_size)
        self.font_size_spin.valueChanged.connect(self._font_size_changed)
        tb_fmt.addWidget(self.font_size_spin)

        add_fmt("Color", self._choose_text_color)
        add_fmt("Highlight", self._choose_highlight)
        tb_fmt.addSeparator()
        add_fmt("B", lambda: self._toggle_format('bold'), "Ctrl+B")
        add_fmt("I", lambda: self._toggle_format('italic'), "Ctrl+I")
        add_fmt("U", lambda: self._toggle_format('underline'), "Ctrl+U")
        add_fmt("S", lambda: self._toggle_format('strike'))
        tb_fmt.addSeparator()
        add_fmt("X₂", lambda: self._toggle_format('subscript'), "Ctrl+=")
        add_fmt("X²", lambda: self._toggle_format('superscript'), "Ctrl+Shift++")
        tb_fmt.addSeparator()
        add_fmt("L", lambda: self.editor.setAlignment(Qt.AlignmentFlag.AlignLeft))
        add_fmt("C", lambda: self.editor.setAlignment(Qt.AlignmentFlag.AlignCenter))
        add_fmt("R", lambda: self.editor.setAlignment(Qt.AlignmentFlag.AlignRight))
        add_fmt("J", lambda: self.editor.setAlignment(Qt.AlignmentFlag.AlignJustify))
        tb_fmt.addSeparator()
        
        add_fmt("I+", lambda: self._modify_indent(1), "Tab")
        add_fmt("I-", lambda: self._modify_indent(-1), "Shift+Tab")
        
        # Line Spacing Menu
        self.btn_spacing = QtWidgets.QToolButton()
        self.btn_spacing.setText("Spacing")
        self.btn_spacing.setPopupMode(QtWidgets.QToolButton.ToolButtonPopupMode.InstantPopup)
        self.menu_spacing = QtWidgets.QMenu(self.btn_spacing)
        self.btn_spacing.setMenu(self.menu_spacing)
        
        self.menu_spacing.clear()
        spacing_options = [("0.5", 50), ("0.8", 80), ("1.0", 100), ("1.15", 115), ("1.5", 150), ("2.0", 200)]
        for space, val in spacing_options:
            self.menu_spacing.addAction(space, lambda v=val: self._set_line_spacing(v))
        
        tb_fmt.addWidget(self.btn_spacing)       
        
        add_fmt("Bullet", self._insert_bullet)
        
        self.btn_numbered = QtWidgets.QToolButton()
        self.btn_numbered.setText("Number")
        self.btn_numbered.setPopupMode(QtWidgets.QToolButton.ToolButtonPopupMode.InstantPopup)
        self.menu_numbered = QtWidgets.QMenu(self.btn_numbered)
        self.btn_numbered.setMenu(self.menu_numbered)
        
        # -----------------------------------------------------------------------------------------
        from PySide6.QtWidgets import QWidgetAction, QWidget, QGridLayout, QPushButton

        grid_widget = QWidget()
        grid_layout = QGridLayout(grid_widget)
        grid_layout.setContentsMargins(5, 5, 5, 5)
        grid_layout.setSpacing(4)

        styles = [
            # 1. Standard Numbers
            ("1.", QtGui.QTextListFormat.Style.ListDecimal, "", "."),
            ("1).", QtGui.QTextListFormat.Style.ListDecimal, "", ")."),
            ("(1).", QtGui.QTextListFormat.Style.ListDecimal, "(", ")."),
            
            # 2. Lower Alpha
            ("a.", QtGui.QTextListFormat.Style.ListLowerAlpha, "", "."),
            ("a).", QtGui.QTextListFormat.Style.ListLowerAlpha, "", ")."),
            ("(a).", QtGui.QTextListFormat.Style.ListLowerAlpha, "(", ")."),
            
            # 3. Upper Alpha
            ("A.", QtGui.QTextListFormat.Style.ListUpperAlpha, "", "."),
            ("A).", QtGui.QTextListFormat.Style.ListUpperAlpha, "", ")."),
            ("(A).", QtGui.QTextListFormat.Style.ListUpperAlpha, "(", ")."),
            
            # 4. Lower Roman
            ("i.", QtGui.QTextListFormat.Style.ListLowerRoman, "", "."),
            ("i).", QtGui.QTextListFormat.Style.ListLowerRoman, "", ")."),
            ("(i).", QtGui.QTextListFormat.Style.ListLowerRoman, "(", ")."),
            
            # 5. Upper Roman
            ("I.", QtGui.QTextListFormat.Style.ListUpperRoman, "", "."),
            ("I).", QtGui.QTextListFormat.Style.ListUpperRoman, "", ")."),
            ("(I).", QtGui.QTextListFormat.Style.ListUpperRoman, "(", ")."),
        ]

        for index, (label, style_fmt, prefix, suffix) in enumerate(styles):
            btn = QPushButton(label)
            btn.setFixedWidth(65)
            btn.clicked.connect(lambda checked=False, s=style_fmt, p=prefix, sx=suffix: [
                self._insert_custom_numbered(s, p, sx),
                self.menu_numbered.close()
            ])
            # Distribute items uniformly into 2 structural columns
            row = index // 2
            col = index % 2
            grid_layout.addWidget(btn, row, col)

        action_widget = QWidgetAction(self)
        action_widget.setDefaultWidget(grid_widget)
        self.menu_numbered.clear()
        self.menu_numbered.addAction(action_widget)
        
        tb_fmt.addWidget(self.btn_numbered)
        # -----------------------------------------------------------------------------------------
        
        tb_fmt.addSeparator()
        add_fmt("Table", self._insert_table)
        add_fmt("Img", self._insert_image)
        tb_fmt.addSeparator()

        # Create an expanding spacer widget to push the voice indicator to the right side of tb_file
        spacer = QtWidgets.QWidget()
        spacer.setSizePolicy(QtWidgets.QSizePolicy.Policy.Expanding, QtWidgets.QSizePolicy.Policy.Preferred)
        tb_file.addWidget(spacer) 

        # Add the visual indicator
        self.lbl_speech_indicator = QtWidgets.QLabel("Hold Ctrl+L to speak ")
        self.lbl_speech_indicator.setStyleSheet("font-weight: bold; padding-right: 10px;")
        tb_file.addWidget(self.lbl_speech_indicator)

    def _populate_template_menu(self):
        self.tpl_menu.clear()
        for name, html in self.state.templates.items():
            self.tpl_menu.addAction(name, lambda h=html: self._insert_template_safe(h))
        self.tpl_menu.addSeparator()
        self.tpl_menu.addAction("Manage Advanced Templates...", self._open_template_manager)

    def _open_suggestion_trainer(self):
        dlg = SuggestionsManagerDialog(self.state, self)
        dlg.exec()

    def _open_prediction_manager(self):
        dlg = PredictionManagerDialog(self.state, self)
        dlg.exec()

    def _open_template_manager(self):
        dlg = TemplateManagerDialog(self.state, self.editor, self)
        dlg.exec()

    def _build_menus(self):
        men = self.menuBar()
        filem = men.addMenu("&File")
        filem.addAction(self.act_new)
        filem.addAction(self.act_open)
        filem.addAction(self.act_save)
        filem.addAction(self.act_save_docx)
        filem.addAction(self.act_save_html)
        filem.addAction(self.act_export_pdf)
        filem.addSeparator()
        filem.addAction(self.act_page_setup)
        filem.addAction(self.act_print)
        
        editm = men.addMenu("&Edit")
        editm.addAction(self.act_undo)
        editm.addAction(self.act_redo)
        editm.addSeparator()
        editm.addAction(self.act_find)
        editm.addSeparator()
        editm.addAction(self.act_insert_date)
        editm.addAction(self.act_insert_line) 
        
        viewm = men.addMenu("&View")
        pagem = viewm.addMenu("Page Layout View")
        pagem.addAction(self.act_view_web)
        pagem.addAction(self.act_view_port)
        pagem.addAction(self.act_view_land)
        suggm = viewm.addMenu("Suggestion UI Mode")
        suggm.addAction(self.act_sugg_dock)
        suggm.addAction(self.act_sugg_inline)
        suggm.addAction(self.act_sugg_both)
        
        popup_design_menu = viewm.addMenu("Suggestion Popup Settings")
        
        style_menu = popup_design_menu.addMenu("Visual Theme")
        self.style_grp = QtGui.QActionGroup(self)
        for s_name in ["Default (OS Native)", "Classic", "Modern", "Minimalist", "Neon", "Google (Search Style)"]:
            act = QtGui.QAction(s_name, self, checkable=True)
            if self.sugg_style == s_name: act.setChecked(True)
            act.triggered.connect(lambda checked, name=s_name: self._change_popup_style(name))
            self.style_grp.addAction(act)
            style_menu.addAction(act)
            
        grid_menu = popup_design_menu.addMenu("Grid & Multi-Column Setup")
        act_cols = QtGui.QAction("Set Number of Columns...", self)
        act_cols.triggered.connect(self._change_sugg_columns)
        act_rows = QtGui.QAction("Set Words per Column (Rows)...", self)
        act_rows.triggered.connect(self._change_sugg_rows)
        grid_menu.addAction(act_cols)
        grid_menu.addAction(act_rows)

        space_menu = popup_design_menu.addMenu("Line Spacing (Density)")
        self.space_grp = QtGui.QActionGroup(self)
        for sp_name in ["Compact", "Normal", "Relaxed"]:
            act = QtGui.QAction(sp_name, self, checkable=True)
            if self.sugg_spacing == sp_name: act.setChecked(True)
            act.triggered.connect(lambda checked, name=sp_name: self._change_sugg_spacing(name))
            self.space_grp.addAction(act)
            space_menu.addAction(act)

        popup_design_menu.addSeparator()

        act_font_size = QtGui.QAction("Set Font Size...", self)
        act_font_size.triggered.connect(self._change_popup_font_size)
        popup_design_menu.addAction(act_font_size)

        act_bold = QtGui.QAction("Bold Text", self, checkable=True)
        act_bold.setChecked(self.sugg_bold)
        act_bold.triggered.connect(self._toggle_sugg_bold)
        popup_design_menu.addAction(act_bold)
        
        act_nums = QtGui.QAction("Show Shortcut Numbers", self, checkable=True)
        act_nums.setChecked(self.sugg_show_numbers)
        act_nums.triggered.connect(self._toggle_sugg_numbers)
        popup_design_menu.addAction(act_nums)

        color_menu = popup_design_menu.addMenu("Text Color")
        act_color = QtGui.QAction("Set Custom Text Color...", self)
        act_color.triggered.connect(self._change_sugg_color)
        act_reset_col = QtGui.QAction("Reset to Auto Color", self)
        act_reset_col.triggered.connect(self._reset_sugg_color)
        color_menu.addAction(act_color)
        color_menu.addAction(act_reset_col)
        
                
        viewm.addSeparator()
        viewm.addAction(self.act_word_wrap)

        toolsm = men.addMenu("&Tools")
        toolsm.addAction(self.act_dict)
        toolsm.addAction(self.act_sugg_manager)
        toolsm.addAction(self.act_phrase_manager)
        toolsm.addAction(self.act_pred_manager) 
        toolsm.addSeparator()
        toolsm.addAction(self.act_translit)
        toolsm.addAction(self.act_theme)
        toolsm.addSeparator()
        toolsm.addAction("Set Autosave Interval...", self._set_autosave_interval)

        self.clock_label = QtWidgets.QLabel()
        self.clock_label.setObjectName("TopClock")  
        
        self.clock_label.setAlignment(Qt.AlignmentFlag.AlignVCenter | Qt.AlignmentFlag.AlignRight)
        
        men.setCornerWidget(self.clock_label, Qt.Corner.TopRightCorner)
        self.clock_label.setText(datetime.now().strftime("%H:%M"))

    def _open_phrase_manager(self):
        dlg = PhraseManagerDialog(self.state, self)
        dlg.exec()
        self.ph_list.clear()
        self.ph_list.addItems(self.state.phrases)

    def _build_sidebar(self):
        self.dock = QtWidgets.QDockWidget("Phrases", self)
        self.dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        w = QtWidgets.QWidget()
        v = QtWidgets.QVBoxLayout(w)
        v.setContentsMargins(0, 0, 0, 0) 
        
        self.ph_list = QtWidgets.QListWidget()
        self.ph_list.addItems(self.state.phrases)
        self.ph_list.itemDoubleClicked.connect(lambda i: self.editor.textCursor().insertText(i.text()))
        self.ph_list.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.ph_list.customContextMenuRequested.connect(self._show_phrase_context_menu)
        v.addWidget(self.ph_list)
        
        w.setLayout(v)
        self.dock.setWidget(w)
        self.addDockWidget(Qt.DockWidgetArea.RightDockWidgetArea, self.dock)
        
        # Add the keyboard shortcut to jump to phrases
        self.act_focus_phrases = QtGui.QAction("Focus Phrases", self, shortcut="Alt+P", triggered=self.ph_list.setFocus)
        self.addAction(self.act_focus_phrases)

        # Override the KeyPressEvent for the list widget directly
        def ph_list_key_press(event):
            if event.key() in (Qt.Key.Key_Return, Qt.Key.Key_Enter):
                item = self.ph_list.currentItem()
                if item:
                    self.editor.textCursor().insertText(item.text())
                    self.editor.setFocus() # Jump back to editor
            else:
                QtWidgets.QListWidget.keyPressEvent(self.ph_list, event)
                
        self.ph_list.keyPressEvent = ph_list_key_press

    def _show_phrase_context_menu(self, pos):
        menu = QtWidgets.QMenu(self.ph_list)
        act_add = menu.addAction("Add new phrase...")
        act_edit = menu.addAction("Edit selected phrase")
        act_rem = menu.addAction("Remove selected phrase")
        
        if self.ph_list.currentRow() < 0: 
            act_edit.setEnabled(False)
            act_rem.setEnabled(False)
            
        action = menu.exec(self.ph_list.mapToGlobal(pos))
        
        if action == act_add: self._add_phrase()
        elif action == act_edit: self._edit_phrase()
        elif action == act_rem: self._remove_phrase()

    def _build_suggestions_sidebar(self):
        self.sugg_dock = QtWidgets.QDockWidget("Suggestions (Ctrl+↑/↓, Tab)", self)
        self.sugg_dock.setAllowedAreas(Qt.DockWidgetArea.LeftDockWidgetArea | Qt.DockWidgetArea.RightDockWidgetArea)
        
        self.sugg_dock.setMinimumWidth(150) 
        
        w = QtWidgets.QWidget()
        v = QtWidgets.QVBoxLayout(w)
        v.setContentsMargins(0, 0, 0, 0)
        self.sugg_list = QtWidgets.QListWidget()
        self.sugg_list.setFocusPolicy(Qt.FocusPolicy.NoFocus) 
        self.sugg_list.itemClicked.connect(lambda i: self.editor.insert_suggestion_from_dock(i.data(Qt.ItemDataRole.UserRole)))
        v.addWidget(self.sugg_list)
        w.setLayout(v)
        self.sugg_dock.setWidget(w)
        self.addDockWidget(Qt.DockWidgetArea.LeftDockWidgetArea, self.sugg_dock)
        
        if self.state.suggestion_mode == "Inline": self.sugg_dock.hide()

    @QtCore.Slot(list)
    def _populate_dock_suggestions(self, matches: list):
        self.sugg_list.clear()
        if not matches:
            self.editor.set_has_dock_suggestions(False)
            return
        for word in matches:
            item = QtWidgets.QListWidgetItem(word)
            item.setData(Qt.ItemDataRole.UserRole, word)
            self.sugg_list.addItem(item)
        if self.sugg_list.count() > 0:
            self.sugg_list.setCurrentRow(0)
            self.editor.set_has_dock_suggestions(True)
        else:
            self.editor.set_has_dock_suggestions(False)

    @QtCore.Slot(int)
    def _navigate_dock_suggestion_list(self, step: int):
        c = self.sugg_list.count()
        if c == 0: return
        
        curr = self.sugg_list.currentRow()
        
        new_row = max(0, min(curr + step, c - 1))
        
        self.sugg_list.setCurrentRow(new_row)

    @QtCore.Slot()
    def _insert_selected_dock_suggestion(self):
        row = self.sugg_list.currentRow()
        if row >= 0:
            word = self.sugg_list.item(row).data(Qt.ItemDataRole.UserRole)
            self.editor.insert_suggestion_from_dock(word)

    def _build_statusbar(self):
        self.status = self.statusBar()
        self.lbl_cursor_pos = QtWidgets.QLabel("Ln 1, Col 0")
        self.lbl_counts = QtWidgets.QLabel("0 words • 0 chars")
        self.lbl_doc_time = QtWidgets.QLabel("Doc: 00:00:00")
        self.lbl_app_time = QtWidgets.QLabel("App: 00:00:00")
        self.lbl_translit_badge = QtWidgets.QLabel()
        
        self.lbl_cursor_pos.setStyleSheet("padding-right: 15px;")
        self.lbl_counts.setStyleSheet("padding-right: 15px;")
        self.lbl_doc_time.setStyleSheet("padding-right: 15px;")
        self.lbl_app_time.setStyleSheet("padding-right: 15px;")
        
        self.status.addPermanentWidget(self.lbl_cursor_pos)
        self.status.addPermanentWidget(self.lbl_counts)
        self.status.addPermanentWidget(self.lbl_doc_time)
        self.status.addPermanentWidget(self.lbl_app_time)
        self.status.addPermanentWidget(self.lbl_translit_badge)

    def _update_translit_badge(self, _=None):
        if self.editor._english_mode:
            self.lbl_translit_badge.setText("EN / HINGLISH")
            self.lbl_translit_badge.setStyleSheet("color: white; background-color: #3b82f6; padding: 2px 6px; border-radius: 4px; font-weight: bold;")
        elif self.state.transliteration_enabled:
            self.lbl_translit_badge.setText("TRANSLIT: ON")
            self.lbl_translit_badge.setStyleSheet("color: white; background-color: #10b981; padding: 2px 6px; border-radius: 4px; font-weight: bold;")
        else:
            self.lbl_translit_badge.setText("TRANSLIT: OFF")
            self.lbl_translit_badge.setStyleSheet("color: white; background-color: #ef4444; padding: 2px 6px; border-radius: 4px; font-weight: bold;")

    def _page_setup(self):
        dialog = QPageSetupDialog(self.printer, self)
        if dialog.exec() == QtWidgets.QDialog.DialogCode.Accepted:
            self._update_view_mode()

    def _change_view_mode(self):
        if self.act_view_port.isChecked(): self.state.view_mode = "Portrait"
        elif self.act_view_land.isChecked(): self.state.view_mode = "Landscape"
        else: self.state.view_mode = "Web"
        self._update_view_mode()

    def _apply_theme(self):
        self.setStyleSheet(style.get_main_theme(self.dark_mode_enabled))

    def _update_view_mode(self):
        fmt = self.editor.document().rootFrame().frameFormat()
        
        if self.state.view_mode in ["Portrait", "Landscape"]:
            self.canvas_layout.setContentsMargins(40, 40, 40, 40)
            layout = self.printer.pageLayout()
            margins = layout.marginsPixels(96)
            
            if self.state.view_mode == "Portrait":
                self.page_frame.setFixedWidth(816)
                self.page_frame.setMinimumHeight(1056)
                self.editor.document().setPageSize(QtCore.QSizeF(816, 1056))
            else:
                self.page_frame.setFixedWidth(1056)
                self.page_frame.setMinimumHeight(816)
                self.editor.document().setPageSize(QtCore.QSizeF(1056, 816))
                
            fmt.setTopMargin(margins.top()); fmt.setBottomMargin(margins.bottom())
            fmt.setLeftMargin(margins.left()); fmt.setRightMargin(margins.right())
            
            if self.dark_mode_enabled:
                fmt.setBackground(QtGui.QColor("#1f2937"))
            else:
                fmt.setBackground(QtGui.QColor("#ffffff"))
                
        else:
            self.page_frame.setMaximumWidth(16777215)
            self.page_frame.setMinimumHeight(0)
            self.canvas_layout.setContentsMargins(0, 0, 0, 0)
            self.editor.document().setPageSize(QtCore.QSizeF(-1, -1))
            fmt.setTopMargin(8); fmt.setBottomMargin(8)
            fmt.setLeftMargin(8); fmt.setRightMargin(8)
            
            if self.dark_mode_enabled:
                fmt.setBackground(QtGui.QColor("#111827"))
            else:
                fmt.clearBackground()
                
        self.editor.document().rootFrame().setFormat(fmt)

    def _toggle_theme(self):
        self.dark_mode_enabled = not self.dark_mode_enabled
        self.settings.setValue("dark_mode", self.dark_mode_enabled)
        self._apply_theme()
        self._apply_popup_style()
        self._update_view_mode()

    def _insert_template_safe(self, html: str):
        self.editor._commit_composing()
        self.editor.insertHtml(html)

    def _toggle_format(self, fmt_type):
        cur = self.editor.textCursor()
        fmt = cur.charFormat()
        if fmt_type == 'bold': fmt.setFontWeight(QtGui.QFont.Weight.Bold if fmt.fontWeight() != QtGui.QFont.Weight.Bold else QtGui.QFont.Weight.Normal)
        elif fmt_type == 'italic': fmt.setFontItalic(not fmt.fontItalic())
        elif fmt_type == 'underline': fmt.setFontUnderline(not fmt.fontUnderline())
        elif fmt_type == 'strike': fmt.setFontStrikeOut(not fmt.fontStrikeOut())
        elif fmt_type == 'subscript':
            fmt.setVerticalAlignment(QTextCharFormat.VerticalAlignment.AlignNormal if fmt.verticalAlignment() == QTextCharFormat.VerticalAlignment.AlignSubScript else QTextCharFormat.VerticalAlignment.AlignSubScript)
        elif fmt_type == 'superscript':
            fmt.setVerticalAlignment(QTextCharFormat.VerticalAlignment.AlignNormal if fmt.verticalAlignment() == QTextCharFormat.VerticalAlignment.AlignSuperScript else QTextCharFormat.VerticalAlignment.AlignSuperScript)
        if cur.hasSelection(): cur.mergeCharFormat(fmt)
        else: self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()

    def _font_changed(self, qfont: QtGui.QFont):
        self.state.font_family = qfont.family()
        self.editor.setFontFamily(qfont.family()) 
        fmt = QtGui.QTextCharFormat()
        fmt.setFont(QtGui.QFont(self.state.font_family, self.state.font_size))
        cur = self.editor.textCursor()
        if cur.hasSelection(): cur.mergeCharFormat(fmt)
        else: self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()

    def _font_size_changed(self, size: int):
        self.state.font_size = size
        self.editor.setFontPointSize(size) 
        fmt = QtGui.QTextCharFormat()
        fmt.setFontPointSize(size)
        cur = self.editor.textCursor()
        if cur.hasSelection(): cur.mergeCharFormat(fmt)
        else: self.editor.mergeCurrentCharFormat(fmt)
        self.editor.setFocus()

    def _choose_text_color(self):
        col = QtWidgets.QColorDialog.getColor(parent=self)
        if col.isValid():
            fmt = QtGui.QTextCharFormat()
            fmt.setForeground(col)
            cur = self.editor.textCursor()
            if cur.hasSelection(): cur.mergeCharFormat(fmt)
            else: self.editor.mergeCurrentCharFormat(fmt)
            self.editor.setFocus()

    def _choose_highlight(self):
        col = QtWidgets.QColorDialog.getColor(parent=self)
        if col.isValid():
            fmt = QtGui.QTextCharFormat()
            fmt.setBackground(col)
            cur = self.editor.textCursor()
            if cur.hasSelection(): cur.mergeCharFormat(fmt)
            else: self.editor.mergeCurrentCharFormat(fmt)
            self.editor.setFocus()

    def _insert_bullet(self): self.editor.textCursor().insertList(QtGui.QTextListFormat.Style.ListDisc)

    def _insert_table(self):
        self.editor._commit_composing()
        rows, ok1 = QtWidgets.QInputDialog.getInt(self, "Insert Table", "Rows:", 3, 1, 100)
        if ok1:
            cols, ok2 = QtWidgets.QInputDialog.getInt(self, "Insert Table", "Columns:", 3, 1, 100)
            if ok2:
                fmt = QtGui.QTextTableFormat()
                fmt.setBorder(1); fmt.setCellPadding(4); fmt.setCellSpacing(0)
                self.editor.textCursor().insertTable(rows, cols, fmt)

    def _insert_image(self):
        self.editor._commit_composing()
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Insert Image", "", "Images (*.png *.jpg *.bmp *.gif *.jpeg)")
        if not path: return
        width, ok = QtWidgets.QInputDialog.getInt(self, "Image Size", "Width (pixels):", 400, 50, 1000)
        if not ok: return
        align_str, ok2 = QtWidgets.QInputDialog.getItem(self, "Alignment", "Choose Alignment:", ["Left", "Center", "Right"], 1, False)
        if not ok2: return

        cur = self.editor.textCursor()
        try:
            img = QtGui.QImage(path)
            if img.isNull(): raise ValueError("Cannot load image")
            img = img.scaledToWidth(width, Qt.TransformationMode.SmoothTransformation)
            fmt = QtGui.QTextImageFormat()
            url = QtCore.QUrl.fromLocalFile(os.path.abspath(path)).toString()
            fmt.setName(url)
            fmt.setWidth(width)
            fmt.setHeight(img.height())
            bf = cur.blockFormat()
            if align_str == "Center": bf.setAlignment(Qt.AlignmentFlag.AlignCenter)
            elif align_str == "Right": bf.setAlignment(Qt.AlignmentFlag.AlignRight)
            else: bf.setAlignment(Qt.AlignmentFlag.AlignLeft)
            cur.setBlockFormat(bf)
            cur.insertImage(fmt)
            self.status.showMessage(f"Inserted image ({width}px, {align_str})", 2000)
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Insert image failed", str(e))

    def _set_autosave_interval(self):
        current_sec = self.autosave_timer.interval() // 1000
        if current_sec == 0: current_sec = 60
        sec, ok = QtWidgets.QInputDialog.getInt(self, "Autosave Interval", "Interval in seconds:", current_sec, 5, 3600)
        if ok:
            self.autosave_timer.setInterval(sec * 1000)
            if self.autosave_enabled:
                self.status.showMessage(f"Autosave interval updated to {sec} seconds", 3000)
                
    def _insert_datetime(self):
        self.editor.textCursor().insertText(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        
    def _insert_horizontal_line(self):
        self.editor._commit_composing()
        self.editor.insertHtml(U_LINE)
        self.editor.setFocus()
        
    def _toggle_wrap(self):
        if self.editor.lineWrapMode() == QtWidgets.QTextEdit.LineWrapMode.NoWrap:
            self.editor.setLineWrapMode(QtWidgets.QTextEdit.LineWrapMode.WidgetWidth)
            self.status.showMessage("Word Wrap: ON", 2000)
        else:
            self.editor.setLineWrapMode(QtWidgets.QTextEdit.LineWrapMode.NoWrap)
            self.status.showMessage("Word Wrap: OFF", 2000)

    def _new_file(self, prompt_autosave=True, is_startup=False):
        if not is_startup and (self.is_dirty or self.editor.toPlainText().strip()):
            reply = QtWidgets.QMessageBox.question(
                self, "Save Current Document?",
                "Do you want to save the current document before starting a new one?",
                QtWidgets.QMessageBox.StandardButton.Yes | QtWidgets.QMessageBox.StandardButton.No | QtWidgets.QMessageBox.StandardButton.Cancel
            )
            if reply == QtWidgets.QMessageBox.StandardButton.Yes:
                self._save_manual()
                if self.is_dirty: return 
            elif reply == QtWidgets.QMessageBox.StandardButton.Cancel:
                return

        self.editor._commit_composing()
        self.editor.clear()
        self.current_filepath = None
        self.autosave_enabled = False
        self.autosave_timer.stop()
        self.is_dirty = False
        
        self.doc_seconds = 0
        self.doc_timer_active = False
        self.lbl_doc_time.setText("Doc: 00:00:00")
        
        if prompt_autosave:
            msg = QtWidgets.QMessageBox(self)
            msg.setWindowTitle("Enable Autosave")
            msg.setText("Would you like to save this new document now to enable continuous auto-saving?")
            btn_save = msg.addButton("Save Now", QtWidgets.QMessageBox.ButtonRole.AcceptRole)
            btn_later = msg.addButton("Later", QtWidgets.QMessageBox.ButtonRole.RejectRole)
            msg.exec()
            if msg.clickedButton() == btn_save:
                path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Initialize Document", "", "Word Document (*.docx);;HTML Web Page (*.html);;Text File (*.txt)")
                if path:
                    self.current_filepath = path
                    self.autosave_enabled = True
                    interval = self.autosave_timer.interval()
                    if interval == 0: interval = 60000
                    self.autosave_timer.start(interval)
                    self.status.showMessage(f"Auto-save enabled (Every {interval//1000}s).")

    def _open_file(self):
        self.editor._commit_composing()
        path, _ = QtWidgets.QFileDialog.getOpenFileName(self, "Open", "", "Documents (*.docx *.html *.htm *.txt *.csv *.xlsx *.ods);;All Files (*.*)")
        if not path: return
        ext = path.lower().split('.')[-1]
        self.status.showMessage("Opening document, please wait...")
        QtWidgets.QApplication.processEvents()
        
        try:
            if ext == "docx" and HAS_DOCX:
                doc = docx.Document(path)
                html_fragments = []
                for p in doc.paragraphs:
                    p_html = "<p"
                    if p.alignment and WD_ALIGN_PARAGRAPH:
                        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER: p_html += ' align="center"'
                        elif p.alignment == WD_ALIGN_PARAGRAPH.RIGHT: p_html += ' align="right"'
                        elif p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY: p_html += ' align="justify"'
                    p_html += ">"
                    for run in p.runs:
                        t = run.text.replace("<", "&lt;").replace(">", "&gt;").replace("\n", "<br>")
                        if run.bold: t = f"<b>{t}</b>"
                        if run.italic: t = f"<i>{t}</i>"
                        if run.underline: t = f"<u>{t}</u>"
                        if run.font.strike: t = f"<s>{t}</s>"
                        p_html += t
                    p_html += "</p>"
                    html_fragments.append(p_html)
                    
                for table in doc.tables:
                    html_fragments.append("<table border='1' cellspacing='0' cellpadding='4'>")
                    for row in table.rows:
                        html_fragments.append("<tr>")
                        for cell in row.cells: html_fragments.append(f"<td>{cell.text}</td>")
                        html_fragments.append("</tr>")
                    html_fragments.append("</table><br>")
                self.editor.setHtml("".join(html_fragments))
            elif ext in ["html", "htm"]:
                with open(path, "r", encoding="utf-8") as f: self.editor.setHtml(f.read())
            elif ext in ["csv", "xlsx", "ods"]:
                if ext == "xlsx": df = pd.read_excel(path, engine="openpyxl")
                elif ext == "ods": df = pd.read_excel(path, engine="odf")
                else: df = pd.read_csv(path)
                self.editor.insertHtml(df.to_html(index=False, border=1) + "<br>")
            else:
                with open(path, "r", encoding="utf-8") as f: self.editor.setPlainText(f.read())
            
            self.current_filepath = path
            self.status.showMessage(f"Opened: {path}", 3000)
            
            self.doc_seconds = 0
            self.doc_timer_active = False
            self.lbl_doc_time.setText("Doc: 00:00:00")
            
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Open failed", str(e))

    def _save_manual(self):
        if self.current_filepath:
            ext = self.current_filepath.lower().split('.')[-1]
            try:
                if ext == "docx" and HAS_DOCX: 
                    self._save_docx_threaded(self.current_filepath)
                elif ext in ["html", "htm"]:
                    self._save_html_silent(self.current_filepath)
                    self.is_dirty = False
                    self.status.showMessage("Saved successfully.")
                else:
                    with open(self.current_filepath, "w", encoding="utf-8") as f: 
                        f.write(self.editor.toPlainText())
                    self.is_dirty = False
                    self.status.showMessage("Saved successfully.")
            except PermissionError:
                QtWidgets.QMessageBox.warning(self, "File Locked", f"Cannot save to {self.current_filepath}. Close the document in Word first.")
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Save Error", str(e))
        else:
            self._save_docx()

    def _save_docx(self):
        if not HAS_DOCX:
            QtWidgets.QMessageBox.warning(self, "Missing", "python-docx required")
            return
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save Document", "", "Word Document (*.docx)")
        if path:
            self._save_docx_threaded(path)

    def _save_docx_threaded(self, path: str):
        self.status.showMessage("Starting background DOCX export...")
        
        cloned_doc = self.editor.document().clone()
        
        self.progress_dialog = QtWidgets.QProgressDialog("Exporting DOCX...", "Cancel", 0, 0, self)
        self.progress_dialog.setWindowTitle("Saving")
        self.progress_dialog.setWindowModality(Qt.WindowModality.WindowModal)
        self.progress_dialog.show()

        self.export_thread = DocxExportThread(cloned_doc, path, self)
        self.export_thread.progress.connect(self.progress_dialog.setLabelText)
        self.export_thread.success.connect(self._on_docx_export_success)
        self.export_thread.error.connect(self._on_docx_export_error)
        self.export_thread.finished.connect(self.progress_dialog.accept)
        
        self.export_thread.start()

    def _on_docx_export_success(self, path: str):
        self.current_filepath = path
        self.is_dirty = False
        self.autosave_enabled = True
        if not self.autosave_timer.isActive():
            self.autosave_timer.start(self.autosave_timer.interval() or 60000)
        self.status.showMessage(f"Saved: {path}")

    def _on_docx_export_error(self, error_msg: str):
        QtWidgets.QMessageBox.critical(self, "Save Error", error_msg)

    def _save_html(self):
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Save as HTML", "", "HTML Web Page (*.html *.htm)")
        if path:
            try:
                self._save_html_silent(path)
                self.status.showMessage(f"Exported HTML to: {path}")
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Save Error", str(e))

    def _save_html_silent(self, path):
        html_content = f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
    body {{
        font-family: '{self.state.font_family}', sans-serif;
        font-size: {self.state.font_size}pt;
        line-height: 1.5;
        margin: auto;
        background-color: #fcfcfc;
    }}
    .document {{

    }}
    table {{ width: 100%; border-collapse: collapse; }}
    td {{ padding: 5px; }}
</style>
</head>
<body>
<div class="document">
{self.editor.toHtml()}
</div>
</body>
</html>"""
        tmp_path = Path(path).with_suffix('.tmp')
        with open(tmp_path, "w", encoding="utf-8") as f:
            f.write(html_content)
        os.replace(tmp_path, path)

    def _export_pdf(self):
        path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Export PDF", "", "PDF Document (*.pdf)")
        if path:
            try:
                print_doc = self.editor.document().clone(self)
                
                fmt = print_doc.rootFrame().frameFormat()
                fmt.setBackground(QtGui.QColor("#ffffff"))
                print_doc.rootFrame().setFormat(fmt)

                printer = QPrinter(QPrinter.PrinterMode.HighResolution)
                printer.setPageLayout(self.printer.pageLayout())
                printer.setOutputFormat(QPrinter.OutputFormat.PdfFormat)
                printer.setOutputFileName(path)
                
                print_doc.print_(printer)
                self.status.showMessage(f"Exported to PDF: {path}")
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Export Error", str(e))

    def _print_doc(self):
        title = APP_NAME + " Document"
        if self.current_filepath: title = os.path.basename(self.current_filepath)
        self.printer.setDocName(title)
        
        try:
            print_doc = self.editor.document().clone(self)
            
            fmt = print_doc.rootFrame().frameFormat()
            fmt.setBackground(QtGui.QColor("#ffffff"))
            print_doc.rootFrame().setFormat(fmt)

            preview = QPrintPreviewDialog(self.printer, self)
            preview.setWindowTitle("Print Preview - " + APP_NAME)
            
            preview.resize(950, 600) 
            
            preview.paintRequested.connect(print_doc.print_)
            preview.exec()
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Print Error", str(e))

    def _open_corrections(self):
        dlg = CorrectionsDialog(self.state, editor_ref=self.editor, parent=self)
        dlg.exec()

    def _train_manual(self):
        cur = self.editor.textCursor()
        if not cur.hasSelection():
            QtWidgets.QMessageBox.information(self, "Info", "Select Devanagari text in the editor first.")
            return
        sel = cur.selectedText()
        latin, ok = QtWidgets.QInputDialog.getText(self, "Train mapping", f"Map Latin shortcut for:\n{sel}")
        if ok and latin.strip():
            self.state.user_dict[latin.strip()] = sel
            self.state.dict_is_dirty = True
            self.status.showMessage(f"Mapped {latin.strip()} -> {sel}")

    def _toggle_transliteration(self):
        self.editor._commit_composing()
        self.state.transliteration_enabled = not self.state.transliteration_enabled
        self._update_translit_badge()
        st = "ON" if self.state.transliteration_enabled else "OFF"
        self.status.showMessage(f"Transliteration {st}", 4000)

    def _update_counts(self, w: int, c: int):
        self.lbl_counts.setText(f"{w} words • {c} chars")

    def _update_cursor_position_label(self, line: int, col: int):
        self.lbl_cursor_pos.setText(f"Ln {line}, Col {col}")

    def _ensure_cursor_visible(self):
        self.editor.ensureCursorVisible()

        if self.state.view_mode in ["Portrait", "Landscape"]:
            cursor_rect = self.editor.cursorRect()
            point_on_canvas = self.editor.mapTo(self.canvas_container, cursor_rect.center())
            self.workspace_scroll.ensureVisible(point_on_canvas.x(), point_on_canvas.y(), 50, 50)

    def _add_phrase(self):
        txt, ok = QtWidgets.QInputDialog.getMultiLineText(self, "Add phrase", "Phrase (Devanagari or Latin):")
        if ok and txt.strip():
            self.state.phrases.append(txt)
            self.ph_list.addItem(txt)
            self.state.save_phrases()

    def _remove_phrase(self):
        r = self.ph_list.currentRow()
        if r >= 0:
            self.state.phrases.pop(r)
            self.ph_list.takeItem(r)
            self.state.save_phrases()
            
    def _edit_phrase(self):
        r = self.ph_list.currentRow()
        if r >= 0:
            current_text = self.state.phrases[r]
            txt, ok = QtWidgets.QInputDialog.getMultiLineText(
                self, "Edit phrase", "Edit Phrase (Devanagari or Latin):", current_text
            )
            if ok and txt.strip():
                self.state.phrases[r] = txt
                self.ph_list.item(r).setText(txt)
                self.state.save_phrases()        

class ImagePropertiesDialog(QtWidgets.QDialog):
    def __init__(self, img_path, current_w, current_h, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Advanced Image Configurations")
        self.img_path = img_path
        self.cropped_path = None
        
        self.PX_TO_CM = 0.0264583333
        self.PX_TO_INCH = 0.0104166667
        
        layout = QtWidgets.QVBoxLayout(self)
        
        geo_box = QtWidgets.QGroupBox("Dimensions Matrix")
        geo_layout = QtWidgets.QGridLayout(geo_box)
        
        geo_layout.addWidget(QtWidgets.QLabel("Unit System:"), 0, 0)
        self.unit_combo = QtWidgets.QComboBox()
        self.unit_combo.addItems(["Pixels (px)", "Centimeters (cm)", "Inches (in)"])
        geo_layout.addWidget(self.unit_combo, 0, 1)
        
        geo_layout.addWidget(QtWidgets.QLabel("Width:"), 1, 0)
        self.sb_w = QtWidgets.QDoubleSpinBox()
        self.sb_w.setRange(1, 5000)
        self.sb_w.setValue(current_w if current_w > 0 else 300)
        geo_layout.addWidget(self.sb_w, 1, 1)
        
        geo_layout.addWidget(QtWidgets.QLabel("Height:"), 2, 0)
        self.sb_h = QtWidgets.QDoubleSpinBox()
        self.sb_h.setRange(1, 5000)
        self.sb_h.setValue(current_h if current_h > 0 else 300)
        geo_layout.addWidget(self.sb_h, 2, 1)
        
        layout.addWidget(geo_box)
        
        crop_btn = QtWidgets.QPushButton("Crop 25% Margins Safely")
        crop_btn.clicked.connect(self._execute_canvas_crop)
        layout.addWidget(crop_btn)
        
        buttons = QtWidgets.QDialogButtonBox(QtWidgets.QDialogButtonBox.StandardButton.Ok | QtWidgets.QDialogButtonBox.StandardButton.Cancel)
        buttons.accepted.connect(self.accept)
        buttons.rejected.connect(self.reject)
        layout.addWidget(buttons)
        
        self.unit_combo.currentIndexChanged.connect(self._convert_measurement_units)
        self.last_unit = 0

    def _convert_measurement_units(self, index):
        w_px = self.sb_w.value()
        h_px = self.sb_h.value()
        
        if self.last_unit == 1: 
            w_px /= self.PX_TO_CM
            h_px /= self.PX_TO_CM
        elif self.last_unit == 2: 
            w_px /= self.PX_TO_INCH
            h_px /= self.PX_TO_INCH

        if index == 0: 
            self.sb_w.setValue(w_px)
            self.sb_h.setValue(h_px)
        elif index == 1: 
            self.sb_w.setValue(w_px * self.PX_TO_CM)
            self.sb_h.setValue(h_px * self.PX_TO_CM)
        elif index == 2: 
            self.sb_w.setValue(w_px * self.PX_TO_INCH)
            self.sb_h.setValue(h_px * self.PX_TO_INCH)
            
        self.last_unit = index

    def _execute_canvas_crop(self):
        pixmap = QtGui.QPixmap(self.img_path)
        if not pixmap.isNull():
            target_rect = QtCore.QRect(0, 0, int(pixmap.width() * 0.75), int(pixmap.height() * 0.75))
            cropped = pixmap.copy(target_rect)
            new_filename = self.img_path.replace(".", "_cropped.")
            if cropped.save(new_filename):
                self.cropped_path = new_filename
                QtWidgets.QMessageBox.information(self, "Success", "Image cropped successfully layout.")

    def get_results(self):
        w = self.sb_w.value()
        h = self.sb_h.value()
        if self.last_unit == 1:
            w /= self.PX_TO_CM; h /= self.PX_TO_CM
        elif self.last_unit == 2:
            w /= self.PX_TO_INCH; h /= self.PX_TO_INCH
        return float(w), float(h), self.cropped_path

def main():
    app = QtWidgets.QApplication(sys.argv)
    app.setApplicationName(APP_NAME)
    
    splash_screen = splash.PremiumSplash()
    splash_screen.start_fade_in()
    
    app.processEvents() 
    
    available_families = QtGui.QFontDatabase.families()
    target_font = "Arial"
    for f in ["Nirmala UI", "Mangal", "Noto Sans Devanagari", "Arial Unicode MS"]:
        if f in available_families:
            target_font = f
            break
            
    if target_font: app.setFont(QtGui.QFont(target_font, 10))
    if not HAS_TRANSLIT:
        QtWidgets.QMessageBox.critical(None, "Missing dependency", "Please install indic-transliteration")
        return
        
    win = MainWindow()
    
    QtCore.QTimer.singleShot(2800, splash_screen.start_fade_out)
    QtCore.QTimer.singleShot(3400, win.showMaximized)
    
    sys.exit(app.exec())

if __name__ == "__main__":
    main()